from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify
from datetime import datetime, timedelta
from sqlalchemy.dialects.sqlite import JSON
from werkzeug.utils import secure_filename
import os
from functools import wraps
import uuid
import subprocess
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from sqlalchemy import or_
from flask_sqlalchemy import SQLAlchemy
from flask import send_from_directory
from flask_mail import Mail, Message
import random
import string
from apscheduler.schedulers.background import BackgroundScheduler
from flask_cors import CORS
from groq import Groq
from sqlalchemy import func
import logging
import json
from collections import Counter
from flask_mail import Message
import re
import pytz
from difflib import SequenceMatcher

app = Flask(__name__)

app.secret_key = 'key'

# Cấu hình thời gian session
app.permanent_session_lifetime = timedelta(days=7)  

# Cấu hình thư mục lưu file
UPLOAD_FOLDER = 'static/uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Đảm bảo thư mục upload tồn tại
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Cấu hình SQLite Database
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://root:@localhost/codejourney'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), nullable=False, unique=True)
    email = db.Column(db.String(120), nullable=False, unique=True)
    password = db.Column(db.String(100), nullable=False)
    admin = db.Column(db.Boolean, default=False)
    avatar = db.Column(db.String(200), nullable=True, default='/static/images/anonymous.svg')
    slogan = db.Column(db.Text, nullable=True)
    points = db.Column(db.Integer, default=0)  # Điểm
    purchase_points = db.Column(db.Integer, default=0) 
    problems_solved = db.Column(db.Integer, default=0)
    problems_solved_today = db.Column(db.Integer, default=0)
    points_earned_today = db.Column(db.Integer, default=0)
    last_point_update = db.Column(db.Date, default=datetime.utcnow)
    tasks_completed = db.Column(db.Integer, default=0) 
    completed_tasks = db.Column(JSON, default=[])
    achievements = db.Column(JSON, default=[])
    selected_frame_id = db.Column(db.Integer, db.ForeignKey('store_item.id'), nullable=True)  # Lưu khung viền đã chọn

    selected_frame = db.relationship('StoreItem', foreign_keys=[selected_frame_id])
    streak = db.Column(db.Integer, default=0)  # Chuỗi tiến độ (số ngày liên tục)
    last_active_date = db.Column(db.Date, nullable=True)

class Lesson(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    content = db.Column(db.Text, nullable=False)
    file_paths = db.Column(JSON, nullable=True)  # Lưu danh sách file

class UserStatusIcon(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    icon_id = db.Column(db.Integer, db.ForeignKey('store_item.id'), nullable=False)
    selected = db.Column(db.Boolean, default=False)  # Biểu tượng hiện tại

    user = db.relationship('User', backref=db.backref('status_icons', lazy=True))
    icon = db.relationship('StoreItem', backref=db.backref('used_as_status', lazy=True))

class UserLesson(db.Model):
    __tablename__ = 'user_lessons'
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), primary_key=True)
    lesson_id = db.Column(db.Integer, db.ForeignKey('lesson.id'), primary_key=True)
    completed = db.Column(db.Boolean, default=False)
    completed_only = db.Column(db.Boolean, default=False)
    redo = db.Column(db.Boolean, default=False)  # Thêm cột mới
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)
    complete_day = db.Column(db.Boolean, default=False)

    user = db.relationship('User', backref=db.backref('user_lessons', lazy='dynamic'))
    lesson = db.relationship('Lesson', backref=db.backref('user_lessons', lazy='dynamic'))

class Challenge(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    description = db.Column(db.Text, nullable=False)
    testcases = db.Column(JSON, nullable=True)  # Dạng JSON lưu testcases
    difficulty = db.Column(db.String(50), default='beginner')
    category = db.Column(db.String(100), nullable=True)

class UserChallenge(db.Model):
    __tablename__ = 'user_challenges'
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), primary_key=True)
    challenge_id = db.Column(db.Integer, db.ForeignKey('challenge.id'), primary_key=True)
    solved = db.Column(db.Boolean, default=False)  # Đã giải đúng chưa
    timestamp = db.Column(db.DateTime, default=datetime.utcnow, nullable=True)
    complete_day = db.Column(db.Boolean, default=False)

    user = db.relationship('User', backref=db.backref('user_challenges', lazy='dynamic'))
    challenge = db.relationship('Challenge', backref=db.backref('user_challenges', lazy='dynamic'))

class StoreItem(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)  # Tên vật phẩm
    description = db.Column(db.Text, nullable=False)  # Mô tả vật phẩm
    price = db.Column(db.Integer, nullable=False)  # Giá vật phẩm
    image = db.Column(db.String(200), default="/static/images/thuoc.svg")  # Đường dẫn hình ảnh
    category = db.Column(db.String(50), nullable=False) 

class PurchasedItem(db.Model):
    __tablename__ = 'purchased_items'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    item_id = db.Column(db.Integer, db.ForeignKey('store_item.id'), nullable=False)
    purchase_date = db.Column(db.DateTime, default=datetime.utcnow)
    
    user = db.relationship('User', backref=db.backref('purchased_items', lazy=True))
    item = db.relationship('StoreItem', backref=db.backref('purchased_by', lazy=True))

class PurchasedFrame(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    frame_id = db.Column(db.Integer, db.ForeignKey('store_item.id'), nullable=False)
    purchase_date = db.Column(db.DateTime, default=datetime.utcnow)

    # Quan hệ
    user = db.relationship('User', backref=db.backref('purchased_frames', lazy=True))
    frame = db.relationship('StoreItem', backref=db.backref('purchased_by_users', lazy=True))


class Organization(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False, unique=True)
    description = db.Column(db.Text, nullable=False)
    org_type = db.Column(db.String(50), nullable=False)  # 'public' hoặc 'private'
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    created_by = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    image_path = db.Column(db.String(255))

    creator = db.relationship('User', backref=db.backref('created_organizations', lazy=True))

class UserOrganization(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    organization_id = db.Column(db.Integer, db.ForeignKey('organization.id'), nullable=False)
    role = db.Column(db.String(50), nullable=False)  # 'admin' hoặc 'member'
    joined_at = db.Column(db.DateTime, default=datetime.utcnow)

    user = db.relationship('User', backref=db.backref('organizations', lazy=True))
    organization = db.relationship('Organization', backref=db.backref('members', lazy=True))

class JoinRequest(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    organization_id = db.Column(db.Integer, db.ForeignKey('organization.id'), nullable=False)
    status = db.Column(db.String(20), nullable=False, default='pending')  # 'pending', 'approved', 'rejected'

    user = db.relationship('User', backref=db.backref('join_requests', lazy=True))
    organization = db.relationship('Organization', backref=db.backref('join_requests', lazy=True))

class OrganizationLesson(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    content = db.Column(db.Text, nullable=False)
    organization_id = db.Column(db.Integer, db.ForeignKey('organization.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    file_paths = db.Column(db.JSON, nullable=True)  # input mẫu
    input_data = db.Column(db.Text, nullable=True)  # output mẫu
    output_data = db.Column(db.Text, nullable=True) # lưu json dạng intput:, output:
    folder_testcases = db.Column(db.JSON, nullable=True)

    organization = db.relationship('Organization', backref=db.backref('lessons', lazy=True))

# Cấu trúc bảng lưu mã khôi phục
class PasswordResetToken(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), nullable=False)
    token = db.Column(db.String(6), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class UserActivity(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    activity_type = db.Column(db.String(50), nullable=False)  # lesson_completed, challenge_solved
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

    user = db.relationship('User', backref=db.backref('activities', lazy=True))

@app.route('/')
def home():
    if 'user_id' in session:
        return redirect(url_for('tongquan'))
    return render_template("index.html", current_page="home")

@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'user_id' in session:
        return redirect(url_for('tongquan'))

    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        remember_me = request.form.get('remember') == 'on'

        # Kiểm tra tài khoản trong cơ sở dữ liệu
        user = User.query.filter_by(email=email).first()
        if user and user.password == password:  # Lưu ý: Nên sử dụng hash mật khẩu (ví dụ, bcrypt) thay vì lưu mật khẩu trực tiếp
            # Lưu thông tin người dùng vào session
            session['user_id'] = user.id
            session['admin'] = user.admin

            if remember_me:
                session.permanent = True  # Session kéo dài (mặc định 31 ngày)
            else:
                session.permanent = False  # Session hết hạn khi đóng trình duyệt

            return redirect(url_for('tongquan'))
        
        session.pop('_flashes', None)

        flash('Invalid email or password!', 'warning')
        return redirect(url_for('login'))

    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    session.pop('_flashes', None)
    if 'user_id' in session:
        return redirect(url_for('tongquan'))

    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        confirm_password = request.form.get('confirmpassword')

        # Kiểm tra mật khẩu khớp
        if password != confirm_password:
            flash('Passwords do not match!', 'warning')
            return render_template('login.html', active_tab='Register')
        
        # Kiểm tra username đã tồn tại
        existing_username = User.query.filter_by(username=username).first()
        if existing_username:
            flash('Username already exists!', 'warning')
            return render_template('login.html', active_tab='Register')

        # Kiểm tra email đã tồn tại
        existing_user = User.query.filter_by(email=email).first()
        if existing_user:
            flash('Email already exists!', 'warning')
            return render_template('login.html', active_tab='Register')

        # Thêm người dùng mới
        new_user = User(
            username=username,
            email=email,
            password=password,  # Lưu ý: Nên sử dụng hash mật khẩu
            achievements=[],
            completed_tasks=[]
        )
        db.session.add(new_user)
        db.session.flush()  # Để có thể lấy được new_user.id
        
        # Thêm bản ghi progress mới với overall = 0
        new_progress = UserProgress(
            user_id=new_user.id,
            overall_completion=0.0
        )
        db.session.add(new_progress)
        db.session.commit()

        flash('Account created successfully!', 'success')
        return redirect(url_for('login'))

    return render_template('login.html', active_tab='Register')

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/logout')
@login_required
def logout():
    session.pop('_flashes', None)
    session.pop('user_id', None)
    session.pop('admin', None)
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))

@app.before_request
def check_session():
    if 'user_id' in session and not session.permanent:
        # Xóa session nếu hết hạn hoặc không chọn Remember Me
        user = db.session.get(User, session['user_id'])

        if not user:
            session.pop('user_id', None)
            session.pop('admin', None)

@app.route('/admin')
@login_required
def admin():
    if 'user_id' not in session or not session.get('admin', False):
        return redirect(url_for('login'))
    
    # Đếm tổng số user
    total_users = db.session.query(User).count()

    # Đếm tổng số bài học (bao gồm cả Lesson và OrganizationLesson)
    total_lessons = db.session.query(Lesson).count() + db.session.query(OrganizationLesson).count()

    # Đếm tổng số thử thách
    total_challenges = db.session.query(Challenge).count()

    # Đếm tổng số vật phẩm trong cửa hàng
    total_store_items = db.session.query(StoreItem).count()
    user_id = session['user_id']
    user = User.query.filter_by(id = user_id).first()

    stats = {
        'total_users': total_users,
        'total_lessons': total_lessons,
        'total_challenges': total_challenges,
        'total_store_items': total_store_items
    }

    return render_template("admin.html", current_page="admin", stats = stats, user=user)

@app.route('/learn')
@login_required
def learn():
    user_id = session['user_id']
    lessons = Lesson.query.all()

    # Lấy danh sách bài học đã từng hoàn thành
    completed_only_lesson_ids = [
        user_lesson.lesson_id
        for user_lesson in UserLesson.query.filter_by(user_id=user_id, completed_only=True).all()
    ]

    return render_template(
        "learn.html",
        lessons=lessons,
        completed_lesson_ids=completed_only_lesson_ids,
        current_page="learn"
    )

@app.route('/learn/<int:id>')
@login_required
def lesson_detail(id):
    lesson = Lesson.query.get_or_404(id)
    user_id = session['user_id']
    # Check if the lesson is completed by the user
    lesson_completed = UserLesson.query.filter_by(user_id=user_id, lesson_id=id, completed=True).first() is not None
    return render_template("lesson_detail.html", lesson=lesson, lesson_completed=lesson_completed, current_page="learn")

@app.route('/complete_lesson/<int:lesson_id>', methods=['POST'])
@login_required
def complete_lesson(lesson_id):
    user_id = session['user_id']
    user = User.query.get(user_id)
    lesson = Lesson.query.get_or_404(lesson_id)

    user_lesson = UserLesson.query.filter_by(user_id=user_id, lesson_id=lesson_id).first()
    if user_lesson:
        user_lesson.completed = True
        user_lesson.redo = False
        user_lesson.completed_only = True  # Đánh dấu đã từng hoàn thành
        user_lesson.timestamp = datetime.utcnow()
        user_lesson.complete_day = True
    else:
        user_lesson = UserLesson(
            user_id=user_id,
            lesson_id=lesson_id,
            completed=True,
            completed_only=True,  # Đánh dấu đã từng hoàn thành
            complete_day=True
        )
        db.session.add(user_lesson)

        user.points += 10
        user.purchase_points += 10
    user.points_earned_today += 10
    
    # Cập nhật tiến độ hàng ngày
    save_daily_progress(user)
    
    db.session.commit()
    return redirect(url_for('lesson_detail', id=lesson_id))

@app.route('/reset')
@login_required
def reset_daily_tasks():
    user = User.query.get(session['user_id'])
    if not user.admin:  
        return redirect(url_for('daily_tasks'))

    if user:
        save_daily_progress(user)
        reset_user_tasks(user)
    return redirect(url_for('daily_tasks'))

@app.route('/redo_lesson/<int:lesson_id>', methods=['POST'])
@login_required
def redo_lesson(lesson_id):
    user_id = session['user_id']
    user_lesson = UserLesson.query.filter_by(user_id=user_id, lesson_id=lesson_id).first()

    if user_lesson:
        # Đánh dấu trạng thái đang học lại
        user_lesson.completed = False
        user_lesson.redo = True
        db.session.commit()
    else:
        # Tạo bản ghi mới nếu chưa tồn tại
        user_lesson = UserLesson(user_id=user_id, lesson_id=lesson_id, completed=False, redo=True)
        db.session.add(user_lesson)
        db.session.commit()

    return redirect(url_for('lesson_detail', id=lesson_id))


@app.route('/complete_task_ajax/<task_key>', methods=['POST'])
@login_required
def complete_task_ajax(task_key):
    user = User.query.get(session['user_id'])

    tasks = {
        'complete_challenge': user.problems_solved_today >= 1,
        'earn_20_points': user.points_earned_today >= 20,
        'complete_lesson': user_lessons_completed_today(user.id) >= 1,
    }

    if task_key in tasks and tasks[task_key]:
        if task_key not in user.completed_tasks:
            # Cập nhật trạng thái hoàn thành
            user.completed_tasks = (user.completed_tasks or []) + [task_key]
            user.tasks_completed += 1
            db.session.commit()

            return jsonify({
                'success': 'Nhiệm vụ đã hoàn thành!',
                'tasks_completed': user.tasks_completed
            }), 200
    return jsonify({'error': 'Nhiệm vụ không hợp lệ hoặc chưa hoàn thành'}), 400

@app.route('/admin/add_challenge', methods=['GET', 'POST'])
@login_required
def add_challenge():
    session.pop('_flashes', None)
    if not session.get('admin', False):
        return redirect(url_for('home'))
        
    if request.method == 'POST':
        description = request.form['description']
        testcases = []
        inputs = request.form.getlist('inputs[]')
        outputs = request.form.getlist('outputs[]')
        sources = request.form.getlist('sources[]')
        difficulty = request.form.get('difficulty')
        category = request.form.get('category')  # Lấy category từ form
        
        for input_data, output_data, source in zip(inputs, outputs, sources):
            if input_data.strip() and output_data.strip():
                testcases.append({"input": input_data, "output": output_data, "source": source})
                
        uploaded_files = request.files.getlist('testcase_folder')
        if uploaded_files:
            folder_testcases = {}
            for file in uploaded_files:
                folder_name = os.path.dirname(file.filename)
                if folder_name not in folder_testcases:
                    folder_testcases[folder_name] = []
                folder_testcases[folder_name].append(file)
                
            for folder, files in folder_testcases.items():
                inp_file = next((f for f in files if f.filename.lower().endswith('test.inp')), None)
                out_file = next((f for f in files if f.filename.lower().endswith('test.out')), None)
                
                if inp_file and out_file:
                    input_data = inp_file.read().decode('utf-8').strip()
                    output_data = out_file.read().decode('utf-8').strip()
                    testcases.append({"input": input_data, "output": output_data, "source": "folder"})
                    
        if not testcases:
            flash('Vui lòng thêm ít nhất một test case!', 'warning')
            return redirect(url_for('add_challenge'))
            
        new_challenge = Challenge(
            description=description,
            testcases=testcases,
            difficulty=difficulty,
            category=category or None  # Lưu None nếu không có category
        )
        db.session.add(new_challenge)
        db.session.commit()
        
        flash('Thêm thử thách thành công!', 'success')
        return redirect(url_for('add_challenge'))
        
    challenges = Challenge.query.all()
    # Danh sách category gợi ý
    suggested_categories = [
        'Số học', 'Tìm kiếm nhị phân', 'Quy hoạch động', 'Đồ thị', 'Xâu',
        'Sắp xếp', 'Tìm kiếm', 'Cây',
    ]
    return render_template(
        'add_challenge.html',
        challenges=challenges,
        current_page="admin",
        categories=suggested_categories
    )

@app.route('/run_code/<int:challenge_id>', methods=['POST'])
@login_required
def run_code(challenge_id):
    challenge = Challenge.query.get_or_404(challenge_id)
    code = request.json.get('code')
    user_id = session['user_id']
    user = db.session.get(User, user_id)

    try:
        test_results = []
        for tc in challenge.testcases:
            try:
                with open("temp_code.py", "w") as file:
                    file.write(code)
                result = subprocess.run(
                    ["python", "temp_code.py"],
                    input=tc["input"],
                    text=True,
                    capture_output=True,
                    timeout=1
                )
                output = result.stdout.strip()
                pass_test = output == tc["output"]
                test_results.append({"input": tc["input"], "output": output, "pass": pass_test})

                if not pass_test:
                    return jsonify({
                        "status": "WRONG",
                        "test_results": test_results
                    })
            except subprocess.TimeoutExpired:
                return jsonify({"status": "TIME_LIMIT"})

        # Nếu tất cả đều đúng
        if all(tc["pass"] for tc in test_results):  # Nếu tất cả các test đều đúng
            problem_solved(user, challenge.id)  # Gọi hàm xử lý cộng điểm
            return jsonify({"status": "ACCEPT", "test_results": test_results})

    except Exception as e:
        print(f"Lỗi: {str(e)}")
        return jsonify({"status": "ERROR", "error": str(e)})

@app.route('/run_code_organization/<int:lesson_id>', methods=['POST'])
@login_required
def run_code_organization(lesson_id):
    lesson = OrganizationLesson.query.get_or_404(lesson_id)
    code = request.json.get('code')
    user_id = session['user_id']
    user = db.session.get(User, user_id)

    if not code:
        return jsonify({"status": "ERROR", "error": "Không có mã code được gửi"})

    test_results = []
    temp_file = "temp_code.py"

    try:
        # Lặp qua tất cả test cases
        for tc in lesson.folder_testcases:
            try:
                # Ghi mã vào file tạm
                with open(temp_file, "w", encoding='utf-8') as file:
                    file.write(code)

                # Chạy mã với input từ test case
                result = subprocess.run(
                    ["python", temp_file],
                    input=tc["input"],
                    text=True,
                    capture_output=True,
                    timeout=5
                )

                # Xử lý output
                output = result.stdout.strip()
                expected_output = tc["output"].strip()
                pass_test = output == expected_output

                # Lưu kết quả test case
                test_results.append({
                    "input": tc["input"],
                    "expected_output": expected_output,
                    "output": output,
                    "pass": pass_test
                })

            except subprocess.TimeoutExpired:
                # Xử lý timeout
                test_results.append({
                    "input": tc["input"],
                    "expected_output": tc["output"],
                    "output": "Timeout",
                    "pass": False
                })
                return jsonify({
                    "status": "TIME_LIMIT",
                    "test_results": test_results
                })
            except Exception as e:
                # Xử lý lỗi khác trong test case
                test_results.append({
                    "input": tc["input"],
                    "expected_output": tc["output"],
                    "output": str(e),
                    "pass": False
                })

        # Kiểm tra kết quả tất cả test cases
        all_passed = all(tc["pass"] for tc in test_results)
        
        return jsonify({
            "status": "ACCEPT" if all_passed else "WRONG",
            "test_results": test_results
        })

    except Exception as e:
        return jsonify({
            "status": "ERROR",
            "error": str(e),
            "test_results": test_results
        })

@app.route('/challenge')
@login_required
def challenge():
    user_id = session['user_id']
    
    # Lấy tham số bộ lọc từ query string
    category_filter = request.args.get('category', None)
    
    # Truy vấn thử thách
    query = Challenge.query.order_by(Challenge.id.asc())
    if category_filter:
        query = query.filter_by(category=category_filter)
    
    challenges = query.all()
    
    # Lấy danh sách thử thách đã hoàn thành
    completed_challenge_ids = [
        uc.challenge_id for uc in UserChallenge.query.filter_by(user_id=user_id, solved=True).all()
    ]
    
    # Lấy danh sách category duy nhất
    categories = db.session.query(Challenge.category).distinct().all()
    categories = [cat[0] for cat in categories if cat[0]]  # Loại bỏ None
    
    return render_template(
        'challenge.html',
        challenges=challenges,
        completed_challenge_ids=completed_challenge_ids,
        current_page="challenge",
        categories=categories,
        selected_category=category_filter
    )

@app.route('/challenge/<int:id>')
@login_required
def challenge_detail(id):
    # Lấy thông tin thử thách từ cơ sở dữ liệu
    challenge = Challenge.query.get_or_404(id)
    return render_template('challenge_detail.html', challenge=challenge, current_page="challenge")

@app.route('/admin/delete_challenge/<int:id>', methods=['POST'])
@login_required
def delete_challenge(id):
    challenge = Challenge.query.get_or_404(id)
    db.session.delete(challenge)
    db.session.commit()
    return redirect(url_for('add_challenge'))

@app.route('/admin/edit_challenge/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_challenge(id):
    session.pop('_flashes', None)
    if not session.get('admin', False):
        return redirect(url_for('home'))
        
    challenge = Challenge.query.get_or_404(id)
    old_folder_testcases = [tc for tc in challenge.testcases if tc.get("source") == "folder"]
    
    if request.method == 'POST':
        challenge.description = request.form['description']
        challenge.difficulty = request.form.get('difficulty') or None  # Thêm difficulty
        challenge.category = request.form.get('category') or None
        
        manual_inputs = request.form.getlist('inputs[]')
        manual_outputs = request.form.getlist('outputs[]')
        manual_testcases = [
            {"input": i, "output": o, "source": "manual"}
            for i, o in zip(manual_inputs, manual_outputs)
            if i.strip() and o.strip()
        ]
        
        uploaded_files = request.files.getlist('testcase_folder')
        folder_testcases = []
        
        if uploaded_files:
            files_by_folder = {}
            for file in uploaded_files:
                folder_name = os.path.dirname(file.filename)
                if folder_name not in files_by_folder:
                    files_by_folder[folder_name] = []
                files_by_folder[folder_name].append(file)
                
            for folder, files in files_by_folder.items():
                inp_file = next((f for f in files if f.filename.lower().endswith("test.inp")), None)
                out_file = next((f for f in files if f.filename.lower().endswith("test.out")), None)
                
                if inp_file and out_file:
                    input_data = inp_file.read().decode('utf-8').strip()
                    output_data = out_file.read().decode('utf-8').strip()
                    folder_testcases.append({"input": input_data, "output": output_data, "source": "folder"})
                    
        if not folder_testcases:
            folder_testcases = old_folder_testcases
            
        challenge.testcases = manual_testcases + folder_testcases
        db.session.commit()
        
        flash('Cập nhật thử thách thành công!', 'success')
        return redirect(url_for('add_challenge'))
        
    suggested_categories = [
        'Số học', 'Tìm kiếm nhị phân', 'Quy hoạch động', 'Đồ thị', 'Xâu',
        'Sắp xếp', 'Tìm kiếm', 'Cây',
    ]
    return render_template(
        'edit_challenge.html',
        challenge=challenge,
        current_page="admin",
        categories=suggested_categories
    )

@app.route('/rank')
@login_required
def rank():
    # Truy vấn tất cả người dùng từ database, sắp xếp theo điểm giảm dần
    users = User.query.order_by(User.points.desc()).all()

    # Nếu có ít hơn 3 người, thêm người dùng ảo
    if len(users) < 3:
        # Lấy danh sách người dùng thật hiện có
        top_users = users

        # Tạo người dùng ảo để đủ 3 người
        fake_users_needed = 3 - len(top_users)
        fake_users = [
            {
                "id": f"{len(top_users) + i + 1}",
                "username": f"anonymous",
                "points": 0,  # Điểm giả định giảm dần
                "avatar": "/static/images/anonymous.svg",  # Avatar mặc định
                "check": "fake_user"
            }
            for i in range(fake_users_needed)
        ]

        # Gộp người dùng thật và người dùng ảo
        top_users += fake_users
        other_users = []
    else:
        # Nếu có đủ hoặc hơn 3 người, lấy danh sách 3 người dùng hàng đầu
        top_users = users[:3]
        other_users = users[3:]

    return render_template(
        "rank.html", 
        top_users=top_users, 
        all_users=other_users, 
        current_page="rank"
    )

def user_lessons_completed_today(user_id):
    today = datetime.utcnow().date()
    completed_lessons = UserLesson.query.filter(
        UserLesson.user_id == user_id,
        (UserLesson.completed == True) | (UserLesson.redo == True),  # Kiểm tra cả trạng thái redo
        db.func.date(UserLesson.timestamp) == today
    ).count()
    return completed_lessons

def reset_daily_counters(user):
    if user is None:
        # Nếu user không tồn tại, không thực hiện gì cả
        return
    today = datetime.utcnow().date()
    if user.last_point_update != today:
        user.points_earned_today = 0
        user.problems_solved_today = 0
        user.last_point_update = today
        db.session.commit()

def reset_user_tasks(user):
    user.completed_tasks = []  # Reset danh sách nhiệm vụ hoàn thành
    user.points_earned_today = 0  # Reset điểm kiếm được hôm nay
    user.problems_solved_today = 0  # Reset số bài giải hôm nay
    UserLesson.query.filter_by(user_id=user.id).update({'redo': False, 'completed': False})
    UserLesson.query.filter_by(user_id=user.id).update({'complete_day': False})
    UserChallenge.query.filter_by(user_id=user.id).update({'complete_day': False})
    user.last_point_update = datetime.utcnow().date()
    db.session.commit()

@app.before_request
def reset_daily_data():
    if 'user_id' in session:
        user = db.session.get(User, session['user_id'])

        if user and user.last_point_update != datetime.utcnow().date():
            reset_user_tasks(user)

def problem_solved(user, challenge_id):
    # Kiểm tra xem người dùng đã giải bài này trước đó chưa
    user_challenge = UserChallenge.query.filter_by(user_id=user.id, challenge_id=challenge_id).first()
    
    if user_challenge:
        if not user_challenge.solved:
            # Lần đầu giải đúng
            user.purchase_points += 10
            user.points += 10  # Cộng điểm tổng cho lần đầu
            user_challenge.solved = True  # Đánh dấu đã giải đúng
            user_challenge.complete_day = True
            user_challenge.timestamp = datetime.utcnow()
        elif not user_challenge.complete_day:
            user_challenge.complete_day = True
            user_challenge.timestamp = datetime.utcnow()
    else:
        # Tạo bản ghi mới nếu chưa tồn tại
        user_challenge = UserChallenge(user_id=user.id, challenge_id=challenge_id, solved=True, complete_day=True)
        user.problems_solved += 1
        user.points += 10  # Cộng điểm tổng cho lần đầu
        user.purchase_points += 10

    # Luôn cộng vào điểm hàng ngày
    user.points_earned_today += 5
    user.problems_solved_today += 1

    db.session.add(user_challenge)
    
    # Cập nhật tiến độ hàng ngày
    save_daily_progress(user)
    
    db.session.commit()

def award_points(user, points):
    user.purchase_points += points
    user.points += points
    user.points_earned_today += points
    db.session.commit()


@app.before_request
def reset_user_counters():
    if 'user_id' in session:
        user = db.session.get(User, session['user_id'])
        if user:
            reset_daily_counters(user)
        else:
            session.pop('user_id')  # Xóa session không hợp lệ

@app.route('/daily_tasks')
@login_required
def daily_tasks():
    user = User.query.get(session['user_id'])

    tasks = {
        'complete_challenge': {
            'title': 'Hoàn thành 1 thử thách',
            'current': user.problems_solved_today,
            'total': 1,
            'completed': 'complete_challenge' in (user.completed_tasks or [])
        },
        'earn_20_points': {
            'title': 'Tăng thêm 20 điểm',
            'current': user.points_earned_today,
            'total': 20,
            'completed': 'earn_20_points' in (user.completed_tasks or [])
        },
        'complete_lesson': {
            'title': 'Hoàn thành 1 bài học',
            'current': user_lessons_completed_today(user.id),
            'total': 1,
            'completed': 'complete_lesson' in (user.completed_tasks or [])
        },
    }

    rewards = {
        'sat_1': {'tasks_required': 1, 'title': 'Huy hiệu Sắt I'},
        'sat_2': {'tasks_required': 5, 'title': 'Huy hiệu Sắt II'},
        'sat_3': {'tasks_required': 10, 'title': 'Huy hiệu Sắt III'},
        'thep_1': {'tasks_required': 15, 'title': 'Huy hiệu Thép I'},
        'thep_2': {'tasks_required': 20, 'title': 'Huy hiệu Thép II'},
        'thep_3': {'tasks_required': 30, 'title': 'Huy hiệu Thép III'},
        'dong_1': {'tasks_required': 40, 'title': 'Huy hiệu Đồng I'},
        'dong_2': {'tasks_required': 50, 'title': 'Huy hiệu Đồng II'},
        'dong_3': {'tasks_required': 60, 'title': 'Huy hiệu Đồng III'},
        'bac_1': {'tasks_required': 70, 'title': 'Huy hiệu Bạc I'},
        'bac_2': {'tasks_required': 80, 'title': 'Huy hiệu Bạc II'},
        'bac_3': {'tasks_required': 90, 'title': 'Huy hiệu Bạc III'},
        'vang_1': {'tasks_required': 100, 'title': 'Huy hiệu Vàng I'},
        'vang_2': {'tasks_required': 120, 'title': 'Huy hiệu Vàng II'},
        'vang_3': {'tasks_required': 140, 'title': 'Huy hiệu Vàng III'},
        'bachkim_1': {'tasks_required': 160, 'title': 'Huy hiệu Bạch Kim I'},
        'bachkim_2': {'tasks_required': 180, 'title': 'Huy hiệu Bạch Kim II'},
        'bachkim_3': {'tasks_required': 200, 'title': 'Huy hiệu Bạch Kim III'},
        'lucbao_1': {'tasks_required': 220, 'title': 'Huy hiệu Lục Bảo I'},
        'lucbao_2': {'tasks_required': 250, 'title': 'Huy hiệu Lục Bảo II'},
        'lucbao_3': {'tasks_required': 280, 'title': 'Huy hiệu Lục Bảo III'},
        'lucbao_4': {'tasks_required': 310, 'title': 'Huy hiệu Lục Bảo IV'},
        'lucbao_5': {'tasks_required': 340, 'title': 'Huy hiệu Lục Bảo V'},
        'kimcuong_1': {'tasks_required': 370, 'title': 'Huy hiệu Kim Cương I'},
        'kimcuong_2': {'tasks_required': 400, 'title': 'Huy hiệu Kim Cương II'},
        'kimcuong_3': {'tasks_required': 430, 'title': 'Huy hiệu Kim Cương III'},
        'kimcuong_4': {'tasks_required': 460, 'title': 'Huy hiệu Kim Cương IV'},
        'kimcuong_5': {'tasks_required': 490, 'title': 'Huy hiệu Kim Cương V'},
        'tinhanh_1': {'tasks_required': 520, 'title': 'Huy hiệu Tinh Anh I'},
        'tinhanh_2': {'tasks_required': 550, 'title': 'Huy hiệu Tinh Anh II'},
        'tinhanh_3': {'tasks_required': 580, 'title': 'Huy hiệu Tinh Anh III'},
        'tinhanh_4': {'tasks_required': 610, 'title': 'Huy hiệu Tinh Anh IV'},
        'tinhanh_5': {'tasks_required': 640, 'title': 'Huy hiệu Tinh Anh V'},
        'caothu_1': {'tasks_required': 700, 'title': 'Huy hiệu Cao Thủ I'},
        'caothu_2': {'tasks_required': 750, 'title': 'Huy hiệu Cao Thủ II'},
        'caothu_3': {'tasks_required': 800, 'title': 'Huy hiệu Cao Thủ III'},
        'caothu_4': {'tasks_required': 850, 'title': 'Huy hiệu Cao Thủ IV'},
        'caothu_5': {'tasks_required': 900, 'title': 'Huy hiệu Cao Thủ V'},
        'caothu_6': {'tasks_required': 950, 'title': 'Huy hiệu Cao Thủ VI'},
        'dinhnoc_1': {'tasks_required': 1000, 'title': 'Đỉnh Nóc Kịch Trần I'},
        'dinhnoc_2': {'tasks_required': 1100, 'title': 'Đỉnh Nóc Kịch Trần II'},
        'dinhnoc_3': {'tasks_required': 1200, 'title': 'Đỉnh Nóc Kịch Trần III'},
        'huyenthoai_1': {'tasks_required': 1300, 'title': 'Danh xưng Huyền Thoại I'},
        'huyenthoai_2': {'tasks_required': 1400, 'title': 'Danh xưng Huyền Thoại II'},
        'huyenthoai_3': {'tasks_required': 1500, 'title': 'Danh xưng Huyền Thoại III'},
        'bavuong_1': {'tasks_required': 1600, 'title': 'Danh xưng Bá Vương I'},
        'bavuong_2': {'tasks_required': 1700, 'title': 'Danh xưng Bá Vương II'},
        'bavuong_3': {'tasks_required': 1800, 'title': 'Danh xưng Bá Vương III'},
        'canke_1': {'tasks_required': 2000, 'title': 'Cận Kề Vô Địch I'},
        'canke_2': {'tasks_required': 2200, 'title': 'Cận Kề Vô Địch II'},
        'canke_3': {'tasks_required': 2400, 'title': 'Cận Kề Vô Địch III'},
        'canke_4': {'tasks_required': 2600, 'title': 'Cận Kề Vô Địch IV'},
        'canke_5': {'tasks_required': 2800, 'title': 'Cận Kề Vô Địch V'},
        'canke_6': {'tasks_required': 3000, 'title': 'Cận Kề Vô Địch VI'},
        'vodichthienha': {'tasks_required': 4000, 'title': 'Vô Địch Thiên Hạ'},
    }

    # Loại bỏ phần thưởng đã nhận
    rewards = {k: v for k, v in rewards.items() if v['title'] not in user.achievements}

    return render_template(
        "daily_tasks.html", 
        tasks=tasks, 
        rewards=rewards, 
        user=user, 
        current_page="daily_tasks"
    )

@app.route('/claim_reward', methods=['POST'])
@login_required
def claim_reward():
    user = User.query.get(session['user_id'])
    reward_key = request.json.get('reward_key')

    rewards = {
        'sat_1': {'tasks_required': 1, 'title': 'Huy hiệu Sắt I'},
        'sat_2': {'tasks_required': 5, 'title': 'Huy hiệu Sắt II'},
        'sat_3': {'tasks_required': 10, 'title': 'Huy hiệu Sắt III'},
        'thep_1': {'tasks_required': 15, 'title': 'Huy hiệu Thép I'},
        'thep_2': {'tasks_required': 20, 'title': 'Huy hiệu Thép II'},
        'thep_3': {'tasks_required': 30, 'title': 'Huy hiệu Thép III'},
        'dong_1': {'tasks_required': 40, 'title': 'Huy hiệu Đồng I'},
        'dong_2': {'tasks_required': 50, 'title': 'Huy hiệu Đồng II'},
        'dong_3': {'tasks_required': 60, 'title': 'Huy hiệu Đồng III'},
        'bac_1': {'tasks_required': 70, 'title': 'Huy hiệu Bạc I'},
        'bac_2': {'tasks_required': 80, 'title': 'Huy hiệu Bạc II'},
        'bac_3': {'tasks_required': 90, 'title': 'Huy hiệu Bạc III'},
        'vang_1': {'tasks_required': 100, 'title': 'Huy hiệu Vàng I'},
        'vang_2': {'tasks_required': 120, 'title': 'Huy hiệu Vàng II'},
        'vang_3': {'tasks_required': 140, 'title': 'Huy hiệu Vàng III'},
        'bachkim_1': {'tasks_required': 160, 'title': 'Huy hiệu Bạch Kim I'},
        'bachkim_2': {'tasks_required': 180, 'title': 'Huy hiệu Bạch Kim II'},
        'bachkim_3': {'tasks_required': 200, 'title': 'Huy hiệu Bạch Kim III'},
        'lucbao_1': {'tasks_required': 220, 'title': 'Huy hiệu Lục Bảo I'},
        'lucbao_2': {'tasks_required': 250, 'title': 'Huy hiệu Lục Bảo II'},
        'lucbao_3': {'tasks_required': 280, 'title': 'Huy hiệu Lục Bảo III'},
        'lucbao_4': {'tasks_required': 310, 'title': 'Huy hiệu Lục Bảo IV'},
        'lucbao_5': {'tasks_required': 340, 'title': 'Huy hiệu Lục Bảo V'},
        'kimcuong_1': {'tasks_required': 370, 'title': 'Huy hiệu Kim Cương I'},
        'kimcuong_2': {'tasks_required': 400, 'title': 'Huy hiệu Kim Cương II'},
        'kimcuong_3': {'tasks_required': 430, 'title': 'Huy hiệu Kim Cương III'},
        'kimcuong_4': {'tasks_required': 460, 'title': 'Huy hiệu Kim Cương IV'},
        'kimcuong_5': {'tasks_required': 490, 'title': 'Huy hiệu Kim Cương V'},
        'tinhanh_1': {'tasks_required': 520, 'title': 'Huy hiệu Tinh Anh I'},
        'tinhanh_2': {'tasks_required': 550, 'title': 'Huy hiệu Tinh Anh II'},
        'tinhanh_3': {'tasks_required': 580, 'title': 'Huy hiệu Tinh Anh III'},
        'tinhanh_4': {'tasks_required': 610, 'title': 'Huy hiệu Tinh Anh IV'},
        'tinhanh_5': {'tasks_required': 640, 'title': 'Huy hiệu Tinh Anh V'},
        'caothu_1': {'tasks_required': 700, 'title': 'Huy hiệu Cao Thủ I'},
        'caothu_2': {'tasks_required': 750, 'title': 'Huy hiệu Cao Thủ II'},
        'caothu_3': {'tasks_required': 800, 'title': 'Huy hiệu Cao Thủ III'},
        'caothu_4': {'tasks_required': 850, 'title': 'Huy hiệu Cao Thủ IV'},
        'caothu_5': {'tasks_required': 900, 'title': 'Huy hiệu Cao Thủ V'},
        'caothu_6': {'tasks_required': 950, 'title': 'Huy hiệu Cao Thủ VI'},
        'dinhnoc_1': {'tasks_required': 1000, 'title': 'Đỉnh Nóc Kịch Trần I'},
        'dinhnoc_2': {'tasks_required': 1100, 'title': 'Đỉnh Nóc Kịch Trần II'},
        'dinhnoc_3': {'tasks_required': 1200, 'title': 'Đỉnh Nóc Kịch Trần III'},
        'huyenthoai_1': {'tasks_required': 1300, 'title': 'Danh xưng Huyền Thoại I'},
        'huyenthoai_2': {'tasks_required': 1400, 'title': 'Danh xưng Huyền Thoại II'},
        'huyenthoai_3': {'tasks_required': 1500, 'title': 'Danh xưng Huyền Thoại III'},
        'bavuong_1': {'tasks_required': 1600, 'title': 'Danh xưng Bá Vương I'},
        'bavuong_2': {'tasks_required': 1700, 'title': 'Danh xưng Bá Vương II'},
        'bavuong_3': {'tasks_required': 1800, 'title': 'Danh xưng Bá Vương III'},
        'canke_1': {'tasks_required': 2000, 'title': 'Cận Kề Vô Địch I'},
        'canke_2': {'tasks_required': 2200, 'title': 'Cận Kề Vô Địch II'},
        'canke_3': {'tasks_required': 2400, 'title': 'Cận Kề Vô Địch III'},
        'canke_4': {'tasks_required': 2600, 'title': 'Cận Kề Vô Địch IV'},
        'canke_5': {'tasks_required': 2800, 'title': 'Cận Kề Vô Địch V'},
        'canke_6': {'tasks_required': 3000, 'title': 'Cận Kề Vô Địch VI'},
        'vodichthienha': {'tasks_required': 4000, 'title': 'Vô Địch Thiên Hạ'},
    }

    if reward_key in rewards:
        reward = rewards[reward_key]
        if user.tasks_completed >= reward['tasks_required']:
            # Kiểm tra và thêm huy hiệu
            if reward['title'] not in (user.achievements or []):
                user.achievements = (user.achievements or []) + [reward['title']]
                db.session.commit()
            return jsonify({'success': f"Bạn đã nhận {reward['title']}!"}), 200

    return jsonify({'error': 'Không thể nhận phần thưởng!'}), 400

@app.route('/claim_points', methods=['POST'])
@login_required
def claim_points():
    user_id = session['user_id']
    user = User.query.get(user_id)
    data = request.get_json()
    task_key = data.get('task_key')

    tasks = {
        'complete_challenge': user.problems_solved_today >= 1,
        'earn_20_points': user.points_earned_today >= 20,
        'complete_lesson': user_lessons_completed_today(user.id) >= 1,
    }

    if task_key in tasks and tasks[task_key]:
        if task_key not in (user.completed_tasks or []):
            user.completed_tasks = (user.completed_tasks or []) + [task_key]
            user.points += 5
            user.purchase_points += 5
            db.session.commit()

            return jsonify({
                'success': True,
                'message': 'Bạn đã nhận được 5 điểm!',
                'tasks_completed': user.tasks_completed,
                'points': user.points
            }), 200

    return jsonify({'success': False, 'error': 'Không thể nhận điểm!'}), 400

@app.route('/store')
@login_required
def store():
    user = User.query.get(session['user_id'])
    icons = StoreItem.query.filter_by(category='icon').all()  # Lấy biểu tượng
    frames = StoreItem.query.filter_by(category='frame').all()  # Lấy khung viền
    purchased_items = PurchasedItem.query.filter_by(user_id=user.id).all()
    purchased_frames = PurchasedFrame.query.filter_by(user_id=user.id).all()

    return render_template(
        "store.html",
        user=user,
        icons=icons,
        frames=frames,
        purchased_items=purchased_items,
        purchased_frames=purchased_frames,
        current_page="store"
    )

@app.route('/buy_item/<int:item_id>', methods=['POST'])
@login_required
def buy_item(item_id):
    user = User.query.get(session['user_id'])
    item = StoreItem.query.get_or_404(item_id)

    # Kiểm tra xem người dùng đã mua vật phẩm này chưa
    if item.category == 'icon':
        existing_purchase = PurchasedItem.query.filter_by(user_id=user.id, item_id=item.id).first()
    elif item.category == 'frame':
        existing_purchase = PurchasedFrame.query.filter_by(user_id=user.id, frame_id=item.id).first()
    else:
        return jsonify({'success': False, 'message': 'Loại vật phẩm không hợp lệ.'}), 400

    if existing_purchase:
        return jsonify({'success': False, 'message': 'Bạn đã mua vật phẩm này rồi.'}), 400

    if user.purchase_points >= item.price:  # Kiểm tra điểm mua vật phẩm
        # Trừ điểm mua vật phẩm
        user.purchase_points -= item.price

        # Thêm vật phẩm vào danh sách đã mua
        if item.category == 'icon':
            new_purchase = PurchasedItem(user_id=user.id, item_id=item.id)
            db.session.add(new_purchase)

            # Thêm vào bảng UserStatusIcon
            new_status_icon = UserStatusIcon(user_id=user.id, icon_id=item.id, selected=False)
            db.session.add(new_status_icon)
        elif item.category == 'frame':
            new_purchase = PurchasedFrame(user_id=user.id, frame_id=item.id)
            db.session.add(new_purchase)

        db.session.commit()

        return jsonify({'success': True, 'message': 'Mua thành công!', 'new_points': user.purchase_points})
    else:
        return jsonify({'success': False, 'message': 'Không đủ điểm để mua vật phẩm.'}), 400


@app.route('/set_status_icon/<int:icon_id>', methods=['POST'])
@login_required
def set_status_icon(icon_id):
    user_id = session['user_id']

    # Kiểm tra xem biểu tượng này đã được mua chưa
    status_icon = UserStatusIcon.query.filter_by(user_id=user_id, icon_id=icon_id).first()

    if status_icon:
        # Đặt tất cả biểu tượng khác thành không được chọn
        UserStatusIcon.query.filter_by(user_id=user_id).update({'selected': False})
        db.session.commit()

        # Đặt biểu tượng này thành được chọn
        status_icon.selected = True
        db.session.commit()

    return '', 204  # Trả về trạng thái thành công nhưng không có nội dung

@app.route('/admin/add_store_item', methods=['GET', 'POST'])
@login_required
def add_store_item():
    if request.method == 'POST':
        name = request.form['name']
        price = int(request.form['price'])
        description = request.form['description']
        category = request.form['category']  # Lấy loại vật phẩm
        image = request.files['image']
        image_filename = 'default_item.png'

        # Lưu ảnh vào thư mục tương ứng
        if image and image.filename != '':
            image_filename = image.filename
            upload_folder = os.path.join(app.config['UPLOAD_FOLDER'], f"set_icon" if category == 'icon' else "frames")
            os.makedirs(upload_folder, exist_ok=True)  # Tạo thư mục nếu chưa tồn tại
            image.save(os.path.join(upload_folder, image_filename))

        # Thêm vật phẩm mới
        new_item = StoreItem(
            name=name,
            price=price,
            description=description,
            category=category,
            image=image_filename
        )
        db.session.add(new_item)
        db.session.commit()
        return redirect(url_for('add_store_item'))

    items = StoreItem.query.all()
    return render_template('add_store_item.html', items=items, current_page="admin")

@app.route('/admin/edit_store_item/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_store_item(id):
    item = StoreItem.query.get_or_404(id)
    if request.method == 'POST':
        # Cập nhật thông tin vật phẩm
        item.name = request.form['name']
        item.price = int(request.form['price'])
        item.description = request.form['description']
        item.category = request.form['category']

        # Xử lý ảnh mới nếu có
        image = request.files['image']
        if image and image.filename != '':
            upload_folder = os.path.join(app.config['UPLOAD_FOLDER'], f"set_icon" if item.category == 'icon' else "frames")
            os.makedirs(upload_folder, exist_ok=True)  # Tạo thư mục nếu chưa tồn tại
            image.save(os.path.join(upload_folder, image.filename))
            item.image = image.filename

        db.session.commit()
        return redirect(url_for('add_store_item'))
    
    return render_template('edit_store_item.html', item=item, current_page="admin")


@app.route('/admin/delete_store_item/<int:id>', methods=['POST'])
@login_required
def delete_store_item(id):
    item = StoreItem.query.get_or_404(id)
    db.session.delete(item)
    db.session.commit()
    return redirect(url_for('add_store_item'))

@app.route('/profile', methods=['GET'])
@login_required
def profile():
    user = User.query.get(session['user_id'])

    # Danh sách các huy hiệu sắp xếp theo thứ tự ưu tiên
    rewards = {
        'sat_1': {'tasks_required': 1, 'title': 'Huy hiệu Sắt I'},
        'sat_2': {'tasks_required': 5, 'title': 'Huy hiệu Sắt II'},
        'sat_3': {'tasks_required': 10, 'title': 'Huy hiệu Sắt III'},
        'thep_1': {'tasks_required': 15, 'title': 'Huy hiệu Thép I'},
        'thep_2': {'tasks_required': 20, 'title': 'Huy hiệu Thép II'},
        'thep_3': {'tasks_required': 30, 'title': 'Huy hiệu Thép III'},
        'dong_1': {'tasks_required': 40, 'title': 'Huy hiệu Đồng I'},
        'dong_2': {'tasks_required': 50, 'title': 'Huy hiệu Đồng II'},
        'dong_3': {'tasks_required': 60, 'title': 'Huy hiệu Đồng III'},
        'bac_1': {'tasks_required': 70, 'title': 'Huy hiệu Bạc I'},
        'bac_2': {'tasks_required': 80, 'title': 'Huy hiệu Bạc II'},
        'bac_3': {'tasks_required': 90, 'title': 'Huy hiệu Bạc III'},
        'vang_1': {'tasks_required': 100, 'title': 'Huy hiệu Vàng I'},
        'vang_2': {'tasks_required': 120, 'title': 'Huy hiệu Vàng II'},
        'vang_3': {'tasks_required': 140, 'title': 'Huy hiệu Vàng III'},
        'bachkim_1': {'tasks_required': 160, 'title': 'Huy hiệu Bạch Kim I'},
        'bachkim_2': {'tasks_required': 180, 'title': 'Huy hiệu Bạch Kim II'},
        'bachkim_3': {'tasks_required': 200, 'title': 'Huy hiệu Bạch Kim III'},
        'lucbao_1': {'tasks_required': 220, 'title': 'Huy hiệu Lục Bảo I'},
        'lucbao_2': {'tasks_required': 250, 'title': 'Huy hiệu Lục Bảo II'},
        'lucbao_3': {'tasks_required': 280, 'title': 'Huy hiệu Lục Bảo III'},
        'lucbao_4': {'tasks_required': 310, 'title': 'Huy hiệu Lục Bảo IV'},
        'lucbao_5': {'tasks_required': 340, 'title': 'Huy hiệu Lục Bảo V'},
        'kimcuong_1': {'tasks_required': 370, 'title': 'Huy hiệu Kim Cương I'},
        'kimcuong_2': {'tasks_required': 400, 'title': 'Huy hiệu Kim Cương II'},
        'kimcuong_3': {'tasks_required': 430, 'title': 'Huy hiệu Kim Cương III'},
        'kimcuong_4': {'tasks_required': 460, 'title': 'Huy hiệu Kim Cương IV'},
        'kimcuong_5': {'tasks_required': 490, 'title': 'Huy hiệu Kim Cương V'},
        'tinhanh_1': {'tasks_required': 520, 'title': 'Huy hiệu Tinh Anh I'},
        'tinhanh_2': {'tasks_required': 550, 'title': 'Huy hiệu Tinh Anh II'},
        'tinhanh_3': {'tasks_required': 580, 'title': 'Huy hiệu Tinh Anh III'},
        'tinhanh_4': {'tasks_required': 610, 'title': 'Huy hiệu Tinh Anh IV'},
        'tinhanh_5': {'tasks_required': 640, 'title': 'Huy hiệu Tinh Anh V'},
        'caothu_1': {'tasks_required': 700, 'title': 'Huy hiệu Cao Thủ I'},
        'caothu_2': {'tasks_required': 750, 'title': 'Huy hiệu Cao Thủ II'},
        'caothu_3': {'tasks_required': 800, 'title': 'Huy hiệu Cao Thủ III'},
        'caothu_4': {'tasks_required': 850, 'title': 'Huy hiệu Cao Thủ IV'},
        'caothu_5': {'tasks_required': 900, 'title': 'Huy hiệu Cao Thủ V'},
        'caothu_6': {'tasks_required': 950, 'title': 'Huy hiệu Cao Thủ VI'},
        'dinhnoc_1': {'tasks_required': 1000, 'title': 'Đỉnh Nóc Kịch Trần I'},
        'dinhnoc_2': {'tasks_required': 1100, 'title': 'Đỉnh Nóc Kịch Trần II'},
        'dinhnoc_3': {'tasks_required': 1200, 'title': 'Đỉnh Nóc Kịch Trần III'},
        'huyenthoai_1': {'tasks_required': 1300, 'title': 'Danh xưng Huyền Thoại I'},
        'huyenthoai_2': {'tasks_required': 1400, 'title': 'Danh xưng Huyền Thoại II'},
        'huyenthoai_3': {'tasks_required': 1500, 'title': 'Danh xưng Huyền Thoại III'},
        'bavuong_1': {'tasks_required': 1600, 'title': 'Danh xưng Bá Vương I'},
        'bavuong_2': {'tasks_required': 1700, 'title': 'Danh xưng Bá Vương II'},
        'bavuong_3': {'tasks_required': 1800, 'title': 'Danh xưng Bá Vương III'},
        'canke_1': {'tasks_required': 2000, 'title': 'Cận Kề Vô Địch I'},
        'canke_2': {'tasks_required': 2200, 'title': 'Cận Kề Vô Địch II'},
        'canke_3': {'tasks_required': 2400, 'title': 'Cận Kề Vô Địch III'},
        'canke_4': {'tasks_required': 2600, 'title': 'Cận Kề Vô Địch IV'},
        'canke_5': {'tasks_required': 2800, 'title': 'Cận Kề Vô Địch V'},
        'canke_6': {'tasks_required': 3000, 'title': 'Cận Kề Vô Địch VI'},
        'vodichthienha': {'tasks_required': 4000, 'title': 'Vô Địch Thiên Hạ'},
    }

    # Xác định danh hiệu cao nhất
    highest_achievement = None
    for reward_id, reward in sorted(rewards.items(), key=lambda x: x[1]['tasks_required'], reverse=True):
        if user.tasks_completed >= reward['tasks_required'] and reward['title'] in user.achievements:
            highest_achievement = {
                'id': reward_id,
                'title': reward['title']
            }
            break

    streak_data = {
        'problems_solved': user.problems_solved,
        'tasks_completed': user.tasks_completed,
        'highest_achievement': highest_achievement['title'] if highest_achievement else "Chưa có huy hiệu",
        'streak': user.streak
    }

    status_icons = UserStatusIcon.query.filter_by(user_id=user.id).all()

    # Lấy danh sách tất cả người dùng, sắp xếp theo điểm giảm dần
    all_users = User.query.order_by(User.points.desc()).all()
    
    # Tìm vị trí của người dùng hiện tại trong bảng xếp hạng
    user_rank = next((index + 1 for index, u in enumerate(all_users) if u.id == user.id), None)
    
    # Thêm biểu tượng top nếu người dùng ở top 3
    top_icon = None
    if user_rank == 1:
        top_icon = 'top1.svg'
    elif user_rank == 2:
        top_icon = 'top2.svg'
    elif user_rank == 3:
        top_icon = 'top3.svg'

    return render_template(
        "profile.html", 
        user=user, 
        streak_data=streak_data, 
        status_icons=status_icons,
        highest_achievement=highest_achievement,
        current_page="profile", 
        user_rank=user_rank, 
        top_icon=top_icon,
        is_current_user=True  # Đây là hồ sơ của chính họ
    )

@app.route('/view_profile/<int:user_id>')
@login_required
def view_profile(user_id):
    user = User.query.get_or_404(user_id)
    current_user_id = session.get('user_id')

    # Danh sách các huy hiệu sắp xếp theo thứ tự ưu tiên
    rewards = {
        'sat_1': {'tasks_required': 1, 'title': 'Huy hiệu Sắt I'},
        'sat_2': {'tasks_required': 5, 'title': 'Huy hiệu Sắt II'},
        'sat_3': {'tasks_required': 10, 'title': 'Huy hiệu Sắt III'},
        'thep_1': {'tasks_required': 15, 'title': 'Huy hiệu Thép I'},
        'thep_2': {'tasks_required': 20, 'title': 'Huy hiệu Thép II'},
        'thep_3': {'tasks_required': 30, 'title': 'Huy hiệu Thép III'},
        'dong_1': {'tasks_required': 40, 'title': 'Huy hiệu Đồng I'},
        'dong_2': {'tasks_required': 50, 'title': 'Huy hiệu Đồng II'},
        'dong_3': {'tasks_required': 60, 'title': 'Huy hiệu Đồng III'},
        'bac_1': {'tasks_required': 70, 'title': 'Huy hiệu Bạc I'},
        'bac_2': {'tasks_required': 80, 'title': 'Huy hiệu Bạc II'},
        'bac_3': {'tasks_required': 90, 'title': 'Huy hiệu Bạc III'},
        'vang_1': {'tasks_required': 100, 'title': 'Huy hiệu Vàng I'},
        'vang_2': {'tasks_required': 120, 'title': 'Huy hiệu Vàng II'},
        'vang_3': {'tasks_required': 140, 'title': 'Huy hiệu Vàng III'},
        'bachkim_1': {'tasks_required': 160, 'title': 'Huy hiệu Bạch Kim I'},
        'bachkim_2': {'tasks_required': 180, 'title': 'Huy hiệu Bạch Kim II'},
        'bachkim_3': {'tasks_required': 200, 'title': 'Huy hiệu Bạch Kim III'},
        'lucbao_1': {'tasks_required': 220, 'title': 'Huy hiệu Lục Bảo I'},
        'lucbao_2': {'tasks_required': 250, 'title': 'Huy hiệu Lục Bảo II'},
        'lucbao_3': {'tasks_required': 280, 'title': 'Huy hiệu Lục Bảo III'},
        'lucbao_4': {'tasks_required': 310, 'title': 'Huy hiệu Lục Bảo IV'},
        'lucbao_5': {'tasks_required': 340, 'title': 'Huy hiệu Lục Bảo V'},
        'kimcuong_1': {'tasks_required': 370, 'title': 'Huy hiệu Kim Cương I'},
        'kimcuong_2': {'tasks_required': 400, 'title': 'Huy hiệu Kim Cương II'},
        'kimcuong_3': {'tasks_required': 430, 'title': 'Huy hiệu Kim Cương III'},
        'kimcuong_4': {'tasks_required': 460, 'title': 'Huy hiệu Kim Cương IV'},
        'kimcuong_5': {'tasks_required': 490, 'title': 'Huy hiệu Kim Cương V'},
        'tinhanh_1': {'tasks_required': 520, 'title': 'Huy hiệu Tinh Anh I'},
        'tinhanh_2': {'tasks_required': 550, 'title': 'Huy hiệu Tinh Anh II'},
        'tinhanh_3': {'tasks_required': 580, 'title': 'Huy hiệu Tinh Anh III'},
        'tinhanh_4': {'tasks_required': 610, 'title': 'Huy hiệu Tinh Anh IV'},
        'tinhanh_5': {'tasks_required': 640, 'title': 'Huy hiệu Tinh Anh V'},
        'caothu_1': {'tasks_required': 700, 'title': 'Huy hiệu Cao Thủ I'},
        'caothu_2': {'tasks_required': 750, 'title': 'Huy hiệu Cao Thủ II'},
        'caothu_3': {'tasks_required': 800, 'title': 'Huy hiệu Cao Thủ III'},
        'caothu_4': {'tasks_required': 850, 'title': 'Huy hiệu Cao Thủ IV'},
        'caothu_5': {'tasks_required': 900, 'title': 'Huy hiệu Cao Thủ V'},
        'caothu_6': {'tasks_required': 950, 'title': 'Huy hiệu Cao Thủ VI'},
        'dinhnoc_1': {'tasks_required': 1000, 'title': 'Đỉnh Nóc Kịch Trần I'},
        'dinhnoc_2': {'tasks_required': 1100, 'title': 'Đỉnh Nóc Kịch Trần II'},
        'dinhnoc_3': {'tasks_required': 1200, 'title': 'Đỉnh Nóc Kịch Trần III'},
        'huyenthoai_1': {'tasks_required': 1300, 'title': 'Danh xưng Huyền Thoại I'},
        'huyenthoai_2': {'tasks_required': 1400, 'title': 'Danh xưng Huyền Thoại II'},
        'huyenthoai_3': {'tasks_required': 1500, 'title': 'Danh xưng Huyền Thoại III'},
        'bavuong_1': {'tasks_required': 1600, 'title': 'Danh xưng Bá Vương I'},
        'bavuong_2': {'tasks_required': 1700, 'title': 'Danh xưng Bá Vương II'},
        'bavuong_3': {'tasks_required': 1800, 'title': 'Danh xưng Bá Vương III'},
        'canke_1': {'tasks_required': 2000, 'title': 'Cận Kề Vô Địch I'},
        'canke_2': {'tasks_required': 2200, 'title': 'Cận Kề Vô Địch II'},
        'canke_3': {'tasks_required': 2400, 'title': 'Cận Kề Vô Địch III'},
        'canke_4': {'tasks_required': 2600, 'title': 'Cận Kề Vô Địch IV'},
        'canke_5': {'tasks_required': 2800, 'title': 'Cận Kề Vô Địch V'},
        'canke_6': {'tasks_required': 3000, 'title': 'Cận Kề Vô Địch VI'},
        'vodichthienha': {'tasks_required': 4000, 'title': 'Vô Địch Thiên Hạ'},
    }

    # Xác định danh hiệu cao nhất
    highest_achievement = None
    for reward_id, reward in sorted(rewards.items(), key=lambda x: x[1]['tasks_required'], reverse=True):
        if user.tasks_completed >= reward['tasks_required'] and reward['title'] in user.achievements:
            highest_achievement = {
                'id': reward_id,
                'title': reward['title']
            }
            break

    streak_data = {
        'problems_solved': user.problems_solved,
        'tasks_completed': user.tasks_completed,
        'highest_achievement': highest_achievement['title'] if highest_achievement else "Chưa có huy hiệu",
        'streak': user.streak
    }

    user_avatar = user.avatar if user.avatar else '/static/images/anonymous.svg'

    # Lấy danh sách tất cả người dùng, sắp xếp theo điểm giảm dần
    all_users = User.query.order_by(User.points.desc()).all()
    
    # Tìm vị trí của người dùng hiện tại trong bảng xếp hạng
    user_rank = next((index + 1 for index, u in enumerate(all_users) if u.id == user.id), None)
    
    # Thêm biểu tượng top nếu người dùng ở top 3
    top_icon = None
    if user_rank == 1:
        top_icon = 'top1.svg'
    elif user_rank == 2:
        top_icon = 'top2.svg'
    elif user_rank == 3:
        top_icon = 'top3.svg'

    return render_template(
        "profile.html", 
        user=user, 
        streak_data=streak_data, 
        user_avatar=user_avatar, 
        user_rank=user_rank, 
        highest_achievement=highest_achievement,
        top_icon=top_icon,
        is_current_user=(user.id == current_user_id), current_page="profile"
    )

@app.route('/edit_profile', methods=['GET', 'POST'])
@login_required
def edit_profile():
    user = User.query.get(session['user_id'])

    if request.method == 'POST':
        # Cập nhật slogan
        user.slogan = request.form.get('slogan')

        # Xử lý ảnh đại diện
        avatar = request.files.get('avatar')
        if avatar and avatar.filename != '':
            avatar_path = os.path.join(app.config['UPLOAD_FOLDER'], avatar.filename).replace("\\", "/")
            avatar.save(avatar_path)
            user.avatar = f"/{avatar_path}" if not avatar_path.startswith("/") else avatar_path

        # Cập nhật khung viền
        selected_frame = request.form.get('selected_frame')
        try:
            user.selected_frame_id = int(selected_frame) if selected_frame and selected_frame.isdigit() else None
        except ValueError:
            user.selected_frame_id = None  # Đặt về None nếu giá trị không hợp lệ# Đặt về mặc định nếu không có lựa chọn

        db.session.commit()
        return redirect(url_for('profile'))

    # Lấy danh sách khung viền đã mua
    frames = PurchasedFrame.query.join(StoreItem).filter(
        PurchasedFrame.user_id == user.id,
        StoreItem.category == 'frame'
    ).with_entities(StoreItem.id, StoreItem.name, StoreItem.image).all()

    return render_template("edit_profile.html", user=user, frames=frames, current_page="profile")

@app.route('/admin/add_lesson', methods=['GET', 'POST'])
@login_required
def add_lesson():
    if request.method == 'POST':
        title = request.form['title']
        content = request.form['content']
        files = request.files.getlist('files')
        file_paths = []

        for file in files:
            if file and file.filename:
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
                file.save(file_path)

                if file.filename.endswith('.docx'):
                    pdf_path = file_path.replace('.docx', '.pdf')
                    convert_docx_to_pdf(file_path, pdf_path)
                    file_paths.append(os.path.basename(pdf_path))
                else:
                    file_paths.append(file.filename)

        new_lesson = Lesson(title=title, content=content, file_paths=file_paths if file_paths else None)
        db.session.add(new_lesson)
        db.session.commit()
        return redirect(url_for('add_lesson'))

    lessons = Lesson.query.all()
    return render_template("add_lesson.html", lessons=lessons, current_page="admin")

@app.route('/admin/delete_lesson/<int:id>', methods=['POST'])
@login_required
def delete_lesson(id):
    lesson = Lesson.query.get_or_404(id)

    # Delete all associated UserLesson records
    UserLesson.query.filter_by(lesson_id=id).delete()

    # Delete the lesson
    db.session.delete(lesson)
    db.session.commit()
    return redirect(url_for('add_lesson'))

@app.route('/admin/edit_lesson/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_lesson(id):
    lesson = Lesson.query.get_or_404(id)

    if request.method == 'POST':
        lesson.title = request.form.get('title')
        lesson.content = request.form.get('content')
        # Lấy danh sách file hiện tại từ database
        current_files = lesson.file_paths or []

        # Xóa file được chọn để loại bỏ
        remove_files = request.form.getlist('remove_files')
        if remove_files:
            current_files = [file for file in current_files if file not in remove_files]

        # Thêm file mới
        files = request.files.getlist('files')
        new_files = []
        for file in files:
            if file and file.filename:
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
                file.save(file_path)

                # Nếu file là .docx, chuyển đổi sang PDF
                if file.filename.endswith('.docx'):
                    pdf_path = file_path.replace('.docx', '.pdf')
                    convert_docx_to_pdf(file_path, pdf_path)
                    new_files.append(os.path.basename(pdf_path))
                else:
                    new_files.append(file.filename)

        # Gộp file cũ và file mới
        all_files = current_files + new_files

        # Cập nhật lại lesson.file_paths
        lesson.file_paths = all_files
        db.session.commit()

        return redirect(url_for('add_lesson'))

    return render_template('edit_lesson.html', lesson=lesson, current_page="admin")

def convert_docx_to_pdf(input_path, output_path):
    """Chuyển đổi file DOCX sang PDF."""
    document = Document(input_path)
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    y_position = height - 50
    c.setFont("Helvetica", 12)

    for paragraph in document.paragraphs:
        if y_position < 50:
            c.showPage()
            c.setFont("Helvetica", 12)
            y_position = height - 50
        c.drawString(50, y_position, paragraph.text)
        y_position -= 15

    c.save()

class OrganizationExam(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    start_time = db.Column(db.DateTime, nullable=False)
    end_time = db.Column(db.DateTime, nullable=False)
    organization_id = db.Column(db.Integer, db.ForeignKey('organization.id'), nullable=False)
    file_path = db.Column(db.String(255), nullable=False)  # Store PDF file paths
    answer_path = db.Column(db.String(255), nullable=False)

    organization = db.relationship('Organization', backref=db.backref('organization_exams', lazy=True))

@app.route('/organization', defaults={'view_mode': 'public'})
@app.route('/organization/<view_mode>')
@login_required
def organization(view_mode):
    user_id = session.get('user_id')

    if view_mode == 'public':
        organizations = Organization.query.filter_by(org_type='public').all()
        toggle_mode = 'private'
        toggle_label = 'Xem Tổ Chức Riêng Tư'
        name = 'Công Khai'
    else:
        organizations = Organization.query.filter_by(org_type='private').all()
        toggle_mode = 'public'
        toggle_label = 'Xem Tổ Chức Công Khai'
        name = 'Riêng Tư'

    # Kiểm tra trạng thái tham gia của từng tổ chức
    user_orgs = UserOrganization.query.filter_by(user_id=user_id).all()
    joined_organizations = {org.organization_id for org in user_orgs}

    return render_template(
        'organization.html',
        organizations=organizations,
        joined_organizations=joined_organizations,
        toggle_mode=toggle_mode,
        toggle_label=toggle_label, current_page="organization",
        name=name 
    )

@app.route('/organization/<int:organization_id>', methods=['GET'])
@login_required
def organization_detail(organization_id):
    # Lấy thông tin tổ chức
    organization = Organization.query.get_or_404(organization_id)
    user_id = session.get('user_id')

    # Kiểm tra vai trò của người dùng trong tổ chức
    user_org = UserOrganization.query.filter_by(user_id=user_id, organization_id=organization_id).first()
    is_member = user_org is not None
    is_admin = user_org and user_org.role == 'admin'
    is_manager = user_org and user_org.role == 'manager'

    # Kiểm tra trạng thái yêu cầu tham gia
    pending_request = JoinRequest.query.filter_by(
        user_id=user_id, organization_id=organization_id, status='pending'
    ).first()

    # Lấy số liệu động
    member_count = UserOrganization.query.filter_by(organization_id=organization_id).count()
    assignment_count = OrganizationLesson.query.filter_by(organization_id=organization_id).count()  # Giả định model Assignment
    exam_count = OrganizationExam.query.filter_by(organization_id=organization_id).count()  # Giả định model Exam

    return render_template(
        'organization_detail.html',
        organization=organization,
        is_member=is_member,
        is_admin=is_admin,
        pending_request=pending_request,
        member_count=member_count,
        assignment_count=assignment_count,
        exam_count=exam_count,
        current_page="organization",
        is_manager=is_manager
    )

@app.route('/organization/<int:organization_id>/promote/<int:user_id>', methods=['POST'])
@login_required
def promote_to_admin(organization_id, user_id):
    organization = Organization.query.get_or_404(organization_id)
    current_user_id = session.get('user_id')

    # Kiểm tra quyền admin
    current_user_org = UserOrganization.query.filter_by(user_id=current_user_id, organization_id=organization_id).first()
    if not current_user_org or current_user_org.role != 'admin':
        return redirect(url_for('manage_members', organization_id=organization_id))

    # Cập nhật vai trò thành admin
    user_org = UserOrganization.query.filter_by(user_id=user_id, organization_id=organization_id).first()
    if user_org:
        user_org.role = 'manager'
        db.session.commit()
    else:
        pass

    return redirect(url_for('manage_members', organization_id=organization_id))

@app.route('/organization/<int:organization_id>/kick/<int:user_id>', methods=['POST'])
@login_required
def kick_member(organization_id, user_id):
    organization = Organization.query.get_or_404(organization_id)
    current_user_id = session.get('user_id')

    # Lấy thông tin user hiện tại và user cần kick
    current_user_org = UserOrganization.query.filter_by(
        user_id=current_user_id, 
        organization_id=organization_id
    ).first()
    user_org = UserOrganization.query.filter_by(
        user_id=user_id, 
        organization_id=organization_id
    ).first()

    # Kiểm tra quyền (admin hoặc manager)
    if not current_user_org or current_user_org.role not in ['admin', 'manager']:
        return redirect(url_for('manage_members', organization_id=organization_id))

    # Không cho phép kick chính mình
    if user_id == current_user_id:
        return redirect(url_for('manage_members', organization_id=organization_id))

    # Admin có thể kick bất kỳ ai trừ admin cuối cùng
    if current_user_org.role == 'admin':
        # Kiểm tra tổng số admin nếu kick admin
        if user_org.role == 'admin':
            total_admins = UserOrganization.query.filter_by(
                organization_id=organization_id, 
                role='admin'
            ).count()
            if total_admins <= 1:
                return redirect(url_for('manage_members', organization_id=organization_id))
    
    # Manager chỉ có thể kick member thường
    elif current_user_org.role == 'manager':
        if user_org.role in ['admin', 'manager']:
            return redirect(url_for('manage_members', organization_id=organization_id))

    # Xóa thành viên nếu điều kiện hợp lệ
    if user_org:
        db.session.delete(user_org)
        db.session.commit()

    return redirect(url_for('manage_members', organization_id=organization_id))

@app.route('/organization/<int:organization_id>/manage_members', methods=['GET', 'POST'])
@login_required
def manage_members(organization_id):
    organization = Organization.query.get_or_404(organization_id)
    current_user_id = session.get('user_id')
    current_user_org = UserOrganization.query.filter_by(organization_id=organization_id).first()

    # Kiểm tra quyền admin
    if not current_user_org or current_user_org.role != 'admin':
        return redirect(url_for('organization_detail', organization_id=organization_id))

    # Lấy danh sách thành viên và vai trò
    members = db.session.query(User, UserOrganization).filter(
        User.id == UserOrganization.user_id,
        UserOrganization.organization_id == organization_id
    ).all()

    # Tính tổng số admin
    total_admins = UserOrganization.query.filter(
        UserOrganization.organization_id == organization_id,
        or_(
            UserOrganization.role == 'admin',
            UserOrganization.role == 'manager'
        )
    ).count()

    # Kiểm tra vai trò của người dùng trong tổ chức
    user_org = UserOrganization.query.filter_by(user_id=current_user_id, organization_id=organization_id).first()
    is_member = user_org is not None
    is_admin = user_org and user_org.role == 'admin'
    is_manager = user_org.role == 'manager'

    return render_template(
        'organization_members.html',
        organization=organization,
        members=members,
        total_admins=total_admins,
        current_user_role=current_user_org.role,
        is_member=is_member,
        is_admin=is_admin,
        is_manager=is_manager,
        current_user_id=current_user_id, current_page="organization"
    )

@app.route('/organization/<int:organization_id>/members')
@login_required
def organization_members(organization_id):
    organization = Organization.query.get_or_404(organization_id)
    current_user_id = session.get('user_id')
    current_user_org = UserOrganization.query.filter_by(organization_id=organization_id).first()

    # Kiểm tra quyền admin
    if not current_user_org:
        return redirect(url_for('organization_detail', organization_id=organization_id))

    # Lấy danh sách thành viên và vai trò
    members = db.session.query(User, UserOrganization).filter(
        User.id == UserOrganization.user_id,
        UserOrganization.organization_id == organization_id
    ).all()

    # Tính tổng số admin
    total_admins = UserOrganization.query.filter(
        UserOrganization.organization_id == organization_id,
        or_(
            UserOrganization.role == 'admin',
            UserOrganization.role == 'manager'
        )
    ).count()

    return render_template(
        'organization_view.html',
        organization=organization,
        members=members,
        total_admins=total_admins,
        current_user_role=current_user_org.role,
        current_user_id=current_user_id, current_page="organization"
    )

@app.route('/my_organizations')
@login_required
def my_organizations():
    user_id = session['user_id']
    # Query organizations created by the user
    organizations = Organization.query.filter_by(created_by=user_id).all()
    
    # Get member counts for all organizations
    member_counts = db.session.query(
        UserOrganization.organization_id,
        func.count(UserOrganization.user_id).label('member_count')
    ).group_by(UserOrganization.organization_id).all()
    
    # Convert member counts to a dictionary for easy lookup
    member_count_dict = {org_id: count for org_id, count in member_counts}
    
    # Add member count to each organization object
    for org in organizations:
        org.member_count = member_count_dict.get(org.id, 0)
    
    return render_template(
        'organization_list.html',
        organizations=organizations,
        title="Tổ Chức Của Tôi",
        current_page="organization"
    )

@app.route('/joined_organizations')
@login_required
def joined_organizations():
    user_id = session['user_id']
    organizations = Organization.query.join(UserOrganization).filter(UserOrganization.user_id == user_id).all()

    member_counts = db.session.query(
        UserOrganization.organization_id,
        func.count(UserOrganization.user_id).label('member_count')
    ).group_by(UserOrganization.organization_id).all()
    
    # Convert member counts to a dictionary for easy lookup
    member_count_dict = {org_id: count for org_id, count in member_counts}
    
    # Add member count to each organization object
    for org in organizations:
        org.member_count = member_count_dict.get(org.id, 0)

    return render_template('organization_list.html', organizations=organizations, title="Tổ Chức Đã Tham Gia", current_page="organization")

@app.route('/join_organization/<int:organization_id>', methods=['POST'])
@login_required
def join_organization(organization_id):
    user_id = session['user_id']
    organization = Organization.query.get_or_404(organization_id)

    # Kiểm tra xem người dùng đã là thành viên hoặc đã có yêu cầu tham gia
    user_org = UserOrganization.query.filter_by(user_id=user_id, organization_id=organization_id).first()
    existing_request = JoinRequest.query.filter_by(user_id=user_id, organization_id=organization_id, status='pending').first()

    if user_org:
        pass
    elif existing_request:
        pass
    else:
        if organization.org_type == 'public':
            # Thêm trực tiếp nếu tổ chức công khai
            new_user_org = UserOrganization(
                user_id=user_id,
                organization_id=organization_id,
                role='member'
            )
            db.session.add(new_user_org)
        elif organization.org_type == 'private':
            # Tạo yêu cầu tham gia cho tổ chức riêng tư
            new_request = JoinRequest(
                user_id=user_id,
                organization_id=organization_id,
                status='pending'
            )
            db.session.add(new_request)

        db.session.commit()

    return redirect(url_for('organization_detail', organization_id=organization_id))


@app.route('/organization/<int:organization_id>/join_requests', methods=['GET', 'POST'])
@login_required
def manage_join_requests(organization_id):
    organization = Organization.query.get_or_404(organization_id)
    user_id = session['user_id']
    
    # Check if user is admin or manager in this organization
    user_org = UserOrganization.query.filter_by(
        user_id=user_id,
        organization_id=organization_id
    ).first()
    
    if not user_org or user_org.role not in ['admin', 'manager']:
        return redirect(url_for('organization_detail', organization_id=organization_id))

    if request.method == 'POST':
        request_id = request.form.get('request_id')
        action = request.form.get('action')  # 'approve' hoặc 'reject'
        join_request = JoinRequest.query.get_or_404(request_id)

        if action == 'approve':
            # Thêm người dùng vào tổ chức
            new_user_org = UserOrganization(
                user_id=join_request.user_id,
                organization_id=join_request.organization_id,
                role='member'
            )
            db.session.add(new_user_org)
            db.session.delete(join_request)  # Xóa yêu cầu sau khi duyệt
            db.session.commit()
        elif action == 'reject':
            db.session.delete(join_request)  # Xóa yêu cầu nếu từ chối
            db.session.commit()

    # Lấy danh sách yêu cầu chờ duyệt
    pending_requests = JoinRequest.query.filter_by(organization_id=organization_id, status='pending').all()

    return render_template(
        'manage_join_requests.html',
        organization=organization,
        requests=pending_requests, current_page="organization"
    )

ALLOWED_EXTENSIONS = {'jpg', 'jpeg', 'png', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/create_organization', methods=['GET', 'POST'])
@login_required
def create_organization():
    if request.method == 'POST':
        name = request.form.get('name')
        description = request.form.get('description')
        org_type = request.form.get('type')
        user_id = session['user_id']
        avatar = request.files.get('avatar')

        # Handle avatar upload (cho phép ảnh trùng)
        image_path = None
        if avatar and allowed_file(avatar.filename):
            if avatar.content_length > 5 * 1024 * 1024:
                return render_template('add_organization.html', current_page="organization")
            
            # Tạo tên file độc đáo để tránh ghi đè
            filename = secure_filename(avatar.filename)
            unique_filename = f"{user_id}_{int(datetime.utcnow().timestamp())}_{filename}"
            save_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
            
            # Đảm bảo thư mục tồn tại
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            
            avatar.save(save_path)
            image_path = f"uploads/{unique_filename}"

        try:
            # Create organization
            new_org = Organization(
                name=name,
                description=description,
                org_type=org_type,
                created_by=user_id,
                image_path=image_path  # Có thể là None hoặc đường dẫn trùng
            )
            db.session.add(new_org)
            db.session.commit()

            # Assign admin role to creator
            new_user_org = UserOrganization(
                user_id=user_id,
                organization_id=new_org.id,
                role='admin'
            )
            db.session.add(new_user_org)
            db.session.commit()

            return redirect(url_for('organization'))
        except Exception as e:
            db.session.rollback()
        
        return render_template('add_organization.html', current_page="organization")

    return render_template('add_organization.html', current_page="organization")

@app.route('/organization/<int:organization_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_organization(organization_id):
    organization = Organization.query.get_or_404(organization_id)
    user_id = session['user_id']
    user_org = UserOrganization.query.filter_by(user_id=user_id, organization_id=organization_id).first()

    # Kiểm tra quyền admin
    if not user_org or user_org.role != 'admin':
        return redirect(url_for('organization_detail', organization_id=organization_id))

    if request.method == 'POST':
        # Lấy dữ liệu từ form
        name = request.form.get('name', '').strip()
        description = request.form.get('description', '').strip()
        org_type = request.form.get('type')
        avatar = request.files.get('avatar')

        # Kiểm tra đầu vào
        errors = {}
        if not name:
            errors['name'] = 'Tên tổ chức không được để trống.'
        if org_type not in ['public', 'private']:
            errors['type'] = 'Loại tổ chức không hợp lệ.'

        # Xử lý ảnh đại diện
        image_path = organization.image_path
        if avatar and avatar.filename:
            if avatar.content_length > 5 * 1024 * 1024:
                errors['avatar'] = 'File ảnh không được vượt quá 5MB.'
            elif not allowed_file(avatar.filename):
                errors['avatar'] = 'Định dạng file không hợp lệ.'
            else:
                # Xóa ảnh cũ nếu có
                if image_path and os.path.exists(os.path.join(app.static_folder, image_path)):
                    os.remove(os.path.join(app.static_folder, image_path))
                
                # Lưu ảnh mới
                filename = secure_filename(avatar.filename)
                unique_filename = f"{organization_id}_{int(datetime.utcnow().timestamp())}_{filename}"
                save_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
                os.makedirs(os.path.dirname(save_path), exist_ok=True)
                avatar.save(save_path)
                image_path = f"uploads/{unique_filename}"

        if errors:
            return render_template(
                'edit_organization.html',
                organization=organization,
                error=errors,
                current_page="organization"
            )

        try:
            # Cập nhật tổ chức
            organization.name = name
            organization.description = description
            organization.org_type = org_type
            organization.image_path = image_path
            db.session.commit()
            return redirect(url_for('organization_detail', organization_id=organization_id))
        except Exception as e:
            db.session.rollback()
            errors['general'] = 'Lỗi khi cập nhật tổ chức. Vui lòng thử lại.'
            return render_template(
                'edit_organization.html',
                organization=organization,
                error=errors,
                current_page="organization"
            )

    return render_template(
        'edit_organization.html',
        organization=organization,
        current_page="organization"
    )

@app.route('/organization/<int:organization_id>/delete', methods=['POST'])
@login_required
def delete_organization(organization_id):
    organization = Organization.query.get_or_404(organization_id)
    user_id = session.get('user_id')
    user_org = UserOrganization.query.filter_by(user_id=user_id, organization_id=organization_id).first()

    # Kiểm tra quyền admin
    if not user_org or user_org.role != 'admin':
        return redirect(url_for('organization_detail', organization_id=organization_id))

    # Xóa tổ chức và các liên kết liên quan
    UserOrganization.query.filter_by(organization_id=organization_id).delete()
    JoinRequest.query.filter_by(organization_id=organization_id).delete()
    db.session.delete(organization)
    db.session.commit()

    return redirect(url_for('organization'))

@app.route('/leave_organization/<int:organization_id>', methods=['POST'])
@login_required
def leave_organization(organization_id):
    user_id = session.get('user_id')

    # Tìm liên kết giữa người dùng và tổ chức
    user_org = UserOrganization.query.filter_by(user_id=user_id, organization_id=organization_id).first()

    if user_org:
        # Kiểm tra nếu người dùng là admin cuối cùng
        if user_org.role == 'admin':
            total_admins = UserOrganization.query.filter_by(organization_id=organization_id, role='admin').count()
            if total_admins <= 1:
                return redirect(url_for('organization_detail', organization_id=organization_id))

        # Xóa liên kết giữa người dùng và tổ chức
        db.session.delete(user_org)
        db.session.commit()
    else:
        pass

    return redirect(url_for('organization'))

@app.route('/organization/<int:organization_id>/assignments', methods=['GET'])
@login_required
def organization_lessons(organization_id):
    organization = Organization.query.get_or_404(organization_id)
    lessons = OrganizationLesson.query.filter_by(organization_id=organization_id).all()
    
    # Kiểm tra quyền admin
    user_id = session.get('user_id')
    user_org = UserOrganization.query.filter_by(user_id=user_id, organization_id=organization_id).first()
    is_admin = user_org.role == 'admin' if user_org else False

    return render_template(
        'organization_lessons.html', 
        organization=organization, 
        lessons=lessons, 
        is_admin=is_admin
        , current_page="organization"
    )

@app.route('/organization/lessons/detail/<int:id>', methods=['GET'])
@login_required
def organization_lesson_detail(id):
    lesson = OrganizationLesson.query.get_or_404(id)
    return render_template("organization_lesson_detail.html", lesson=lesson, current_page="organization")

@app.route('/organization/<int:organization_id>/assignments/add', methods=['GET', 'POST'])
@login_required
def add_organization_lesson(organization_id):
    organization = Organization.query.get_or_404(organization_id)

    # Kiểm tra quyền admin của người dùng trong tổ chức
    user_org = UserOrganization.query.filter_by(user_id=session['user_id'], organization_id=organization_id).first()
    if not user_org or user_org.role != 'admin':
        return redirect(url_for('organization_lessons', organization_id=organization_id))

    if request.method == 'POST':
        title = request.form['title']
        content = request.form['content']
        input_data = request.form['input']  # Sample input
        output_data = request.form['output']  # Sample output

        # Xử lý upload tệp tin
        files = request.files.getlist('files')
        file_paths = []
        for file in files:
            if file and file.filename != '':
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
                file.save(file_path)
                file_paths.append(file.filename)

        # Xử lý upload folder test case
        uploaded_files = request.files.getlist('testcase_folder')
        folder_testcases = []
        if uploaded_files:
            folder_testcases_dict = {}
            for file in uploaded_files:
                folder_name = os.path.dirname(file.filename)
                if folder_name not in folder_testcases_dict:
                    folder_testcases_dict[folder_name] = []
                folder_testcases_dict[folder_name].append(file)

            for folder, files in folder_testcases_dict.items():
                inp_file = next((f for f in files if f.filename.lower().endswith('test.inp')), None)
                out_file = next((f for f in files if f.filename.lower().endswith('test.out')), None)

                if inp_file and out_file:
                    input_data_from_file = inp_file.read().decode('utf-8').strip()
                    output_data_from_file = out_file.read().decode('utf-8').strip()
                    folder_testcases.append({
                        "input": input_data_from_file,
                        "output": output_data_from_file
                    })

        print(folder_testcases)

        # Lưu bài học vào cơ sở dữ liệu
        new_lesson = OrganizationLesson(
            title=title,
            content=content,
            organization_id=organization_id,
            created_at=datetime.utcnow(),
            file_paths=file_paths if file_paths else None,
            input_data=input_data if input_data.strip() else None,
            output_data=output_data if output_data.strip() else None,
            folder_testcases=folder_testcases if folder_testcases else None
        )

        db.session.add(new_lesson)
        db.session.commit()
        return redirect(url_for('organization_lessons', organization_id=organization_id))

    return render_template('add_organization_lesson.html', organization=organization, current_page="organization")

@app.route('/organization/<int:organization_id>/assignments/edit/<int:lesson_id>', methods=['GET', 'POST'])
@login_required
def edit_organization_lesson(organization_id, lesson_id):
    lesson = OrganizationLesson.query.get_or_404(lesson_id)
    organization = Organization.query.get_or_404(organization_id)

    # Verify admin access
    user_org = UserOrganization.query.filter_by(user_id=session['user_id'], organization_id=organization_id).first()
    if not user_org or user_org.role != 'admin':
        return redirect(url_for('organization_lessons', organization_id=organization_id))

    if request.method == 'POST':
        # Update title and content
        lesson.title = request.form['title']
        lesson.content = request.form['content']
        lesson.input_data = request.form['input']  # Sample input
        lesson.output_data = request.form['output']  # Sample output

        old_folder_test = lesson.folder_testcases

        # Manage uploaded files (File logic not modified)
        current_files = lesson.file_paths or []
        remove_files = request.form.getlist('remove_files')
        current_files = [file for file in current_files if file not in remove_files]

        new_files = []
        files = request.files.getlist('files')
        for file in files:
            if file and file.filename != '':
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
                file.save(file_path)
                new_files.append(file.filename)

        lesson.file_paths = current_files + new_files

        # Manage folder test cases
        uploaded_files = request.files.getlist('testcase_folder')
        folder_testcases = []

        if uploaded_files:
            folder_testcases_dict = {}
            # Group files by their folder structure
            for file in uploaded_files:
                folder_name = os.path.dirname(file.filename)
                if folder_name not in folder_testcases_dict:
                    folder_testcases_dict[folder_name] = []
                folder_testcases_dict[folder_name].append(file)

            # Read each folder's .inp and .out files and store them as test cases
            for folder, files in folder_testcases_dict.items():
                inp_file = next((f for f in files if f.filename.lower().endswith('test.inp')), None)
                out_file = next((f for f in files if f.filename.lower().endswith('test.out')), None)

                if inp_file and out_file:
                    input_data = inp_file.read().decode('utf-8').strip()
                    output_data = out_file.read().decode('utf-8').strip()
                    folder_testcases.append({
                        "input": input_data,
                        "output": output_data
                    })

        # Update the lesson's folder test cases
        if not folder_testcases:
            lesson.folder_testcases = old_folder_test
        else:
            lesson.folder_testcases = folder_testcases

        # Commit changes to the database
        db.session.commit()
        return redirect(url_for('organization_lessons', organization_id=organization_id))

    return render_template(
        'edit_organization_lesson.html',
        lesson=lesson,
        organization=organization,
        current_page="organization"
    )

@app.route('/organization/<int:organization_id>/assignments/delete/<int:lesson_id>', methods=['POST'])
@login_required
def delete_organization_lesson(organization_id, lesson_id):  # Đổi tên hàm
    lesson = OrganizationLesson.query.get_or_404(lesson_id)

    # Kiểm tra quyền admin
    user_org = UserOrganization.query.filter_by(user_id=session['user_id'], organization_id=organization_id).first()
    if not user_org or user_org.role != 'admin':
        return redirect(url_for('organization_lessons', organization_id=organization_id))

    # Xóa bài học
    db.session.delete(lesson)
    db.session.commit()
    return redirect(url_for('organization_lessons', organization_id=organization_id))

@app.route('/organization/<int:organization_id>/demote/<int:user_id>', methods=['POST'])
@login_required
def demote_admin(organization_id, user_id):
    organization = Organization.query.get_or_404(organization_id)
    current_user_id = session.get('user_id')

    # Kiểm tra xem người gửi yêu cầu có phải là admin tạo ra tổ chức không
    if organization.created_by != current_user_id:
        return redirect(url_for('manage_members', organization_id=organization_id))

    # Tìm admin để cách chức
    user_org = UserOrganization.query.filter_by(user_id=user_id, organization_id=organization_id, role='manager').first()
    if user_org:
        user_org.role = 'member'  # Chuyển vai trò thành thành viên
        db.session.commit()
    else:
        pass

    return redirect(url_for('manage_members', organization_id=organization_id))

class ExamSubject(db.Model):
    __tablename__ = 'exam_subjects'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    subject = db.Column(db.String(50), nullable=False)
    can_upload = db.Column(db.Boolean, default=False)

    user = db.relationship('User', backref=db.backref('exam_subjects', lazy=True))

class SubjectFile(db.Model):
    __tablename__ = 'subject_files'
    id = db.Column(db.Integer, primary_key=True)
    subject = db.Column(db.String(50), nullable=False)  # Môn học
    title = db.Column(db.String(100), nullable=False)  # Tiêu đề file
    file_path = db.Column(db.String(200), nullable=False)  # Đường dẫn file
    uploaded_at = db.Column(db.DateTime, default=datetime.utcnow)  # Thời gian tải lên
    user_id = db.Column(db.Integer, nullable=False)

#######################################################################################################

@app.route('/exam', methods=['GET'])
@login_required
def exam():
    # Danh sách môn học
    subjects = ['Toán', 'Lí', 'Hóa', 'Sinh Học', 'Lịch Sử', 'Địa Lí', 'GDKT-PL', 'Ngữ Văn', 'Tiếng Anh', 'Tin Học', 'Công Nghệ', 'DGNL']

    file_counts = db.session.query(
        SubjectFile.subject,
        func.count(SubjectFile.id).label('count')
    ).group_by(SubjectFile.subject).all()

    # Convert to dictionary for easier template access
    file_counts_dict = {subject: count for subject, count in file_counts}

    # List of all subjects (for consistency with your HTML)
    all_subjects = [
        'Toán', 'Lí', 'Hóa', 'Sinh Học', 'Lịch Sử', 'Địa Lí',
        'GDKT-PL', 'Ngữ Văn', 'Tiếng Anh', 'Tin Học', 'Công Nghệ', 'DGNL'
    ]

    # Ensure every subject has a count (default to 0 if no files)
    subject_counts = {subject: file_counts_dict.get(subject, 0) for subject in all_subjects}
    
    return render_template('exam.html', subjects=subjects, current_page='exam', subject_counts=subject_counts)

@app.route('/view_subject/<string:subject>')
def view_subject(subject):
    user_id = session.get('user_id')  # Lấy user_id từ session
    if not user_id:
        return redirect(url_for('login'))

    # Kiểm tra quyền tải lên của người dùng
    can_upload = ExamSubject.query.filter_by(user_id=user_id, subject=subject, can_upload=True).first() is not None

    # Lấy danh sách file của môn học
    files = SubjectFile.query.filter_by(subject=subject).order_by(SubjectFile.uploaded_at.desc()).all()
    
    return render_template('view_subject.html', subject=subject, files=files, user_can_upload=can_upload, current_page='exam')

@app.route('/upload_file/<string:subject>', methods=['GET', 'POST'])
def upload_file(subject):
    session.pop('_flashes', None)
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))

    # Kiểm tra quyền tải lên
    can_upload = ExamSubject.query.filter_by(user_id=user_id, subject=subject, can_upload=True).first()
    if not can_upload:
        return redirect(url_for('view_subject', subject=subject))

    if request.method == 'POST':
        title = request.form['title']
        uploaded_file = request.files['file']
        if uploaded_file:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], uploaded_file.filename)
            uploaded_file.save(file_path)

            # Lưu thông tin file vào cơ sở dữ liệu, bao gồm user_id
            new_file = SubjectFile(subject=subject, title=title, file_path=uploaded_file.filename, user_id=user_id)
            db.session.add(new_file)
            db.session.commit()
            flash('Tải file lên thành công!', 'success')

        return redirect(url_for('upload_file', subject=subject))

    # Lấy danh sách file hiện có cho môn học
    files = SubjectFile.query.filter_by(subject=subject, user_id=user_id).order_by(SubjectFile.uploaded_at.desc()).all()
    return render_template('upload_file.html', subject=subject, files=files, current_page='exam')

@app.route('/delete_file/<int:file_id>', methods=['POST'])
def delete_file(file_id):
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))

    file = SubjectFile.query.get(file_id)
    if not file:
        return redirect(request.referrer)

    # Kiểm tra xem user hiện tại có phải là người tải file lên không
    if file.user_id != user_id:
        return redirect(request.referrer)

    # Kiểm tra quyền tải lên
    can_upload = ExamSubject.query.filter_by(user_id=user_id, subject=file.subject, can_upload=True).first()
    if not can_upload:
        return redirect(request.referrer)

    # Xóa file khỏi thư mục
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.file_path)
    if os.path.exists(file_path):
        os.remove(file_path)

    # Xóa bản ghi trong cơ sở dữ liệu
    db.session.delete(file)
    db.session.commit()
    return redirect(request.referrer)

@app.route('/download_file/<int:file_id>')
def download_file(file_id):
    # Tìm file trong cơ sở dữ liệu
    file = SubjectFile.query.get(file_id)
    if not file:
        return redirect(request.referrer)

    # Tạo đường dẫn file
    try:
        return send_from_directory(
            directory=app.config['UPLOAD_FOLDER'],
            path=file.file_path,
            as_attachment=True  # Tải file xuống
        )
    except FileNotFoundError:
        return redirect(request.referrer)

#######################################################################################################

@app.route('/check_username', methods=['POST'])
def check_username():
    data = request.get_json()
    username = data.get('username', '').strip()
    
    if not username:
        return jsonify({'exists': False}), 400  # Không nhập gì
    
    user = User.query.filter_by(username=username).first()
    if user:
        return jsonify({'exists': True}), 200
    else:
        return jsonify({'exists': False}), 200


@app.route('/admin/manage_permissions', methods=['GET', 'POST'])
@login_required
def manage_permissions():
    if not session.get('admin'):  # Chỉ admin mới có quyền truy cập
        return redirect(url_for('home'))

    if request.method == 'POST':
        username = request.form['username']
        subject = request.form['subject']
        action = request.form['can_upload']

        # Tìm user bằng tên
        user = User.query.filter_by(username=username).first()
        if not user:
            return redirect(url_for('manage_permissions'))

        # Tìm quyền của user
        permission = ExamSubject.query.filter_by(user_id=user.id, subject=subject).first()
        if action == 'allow':
            # Nếu có quyền thì cập nhật hoặc tạo mới
            if permission:
                permission.can_upload = True
            else:
                new_permission = ExamSubject(user_id=user.id, subject=subject, can_upload=True)
                db.session.add(new_permission)
        elif action == 'deny':
            # Nếu không có quyền thì xóa khỏi database
            if permission:
                db.session.delete(permission)

        db.session.commit()
        return redirect(url_for('manage_permissions'))

    # Lấy danh sách môn học và quyền hiện tại
    subjects = ['Toán', 'Lí', 'Hóa', 'Sinh Học', 'Lịch Sử', 'Địa Lí', 'GDKT-PL', 'Ngữ Văn', 'Tiếng Anh', 'Tin Học', 'Công Nghệ']
    permissions = ExamSubject.query.all()  # Lấy tất cả quyền hiện tại để hiển thị

    return render_template('manage_permissions.html', subjects=subjects, permissions=permissions, current_page='admin')

############################################################'

# Cấu hình Flask-Mail
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'EMAIL_OF_YOU'
app.config['MAIL_PASSWORD'] = 'PASSWORD_OF_YOU'
mail = Mail(app)

# Hàm tạo mã ngẫu nhiên
def generate_reset_code(length=6):
    characters = string.ascii_letters + string.digits
    return ''.join(random.choice(characters) for _ in range(length))

def send_reset_email(email, reset_code):
    msg = Message(
        subject="Password Reset Code",
        sender=app.config['MAIL_USERNAME'],
        recipients=[email]
    )
    msg.body = f"Mã khôi phục mật khẩu của bạn là: {reset_code}"
    mail.send(msg)

def delete_expired_tokens():
    """Xóa tất cả các token đã hết hạn (quá 1 phút)"""
    with app.app_context():  # Đảm bảo chạy trong application context
        expiration_time = datetime.utcnow() - timedelta(minutes=1)
        expired_tokens = PasswordResetToken.query.filter(PasswordResetToken.created_at < expiration_time).all()

        #print(f"Found {len(expired_tokens)} expired tokens.")  # Debugging log
        for token in expired_tokens:
            print(f"Deleting token with ID {token.id} and email {token.email}")
            db.session.delete(token)
        
        db.session.commit()

scheduler = BackgroundScheduler()
with app.app_context():
    scheduler.add_job(delete_expired_tokens, 'interval', minutes=0.1)
    scheduler.start()

@app.before_request
def start_scheduler():
    """Khởi động lịch trình khi ứng dụng bắt đầu"""
    with app.app_context():
        if not scheduler.running:
            scheduler.start()

@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    session.pop('_flashes', None)
    if request.method == 'POST':
        email = request.form['email']

        # Kiểm tra xem email có tồn tại trong bảng User không
        user = User.query.filter_by(email=email).first()
        if not user:
            flash('Account does not exist!', 'warning')
            return redirect(url_for('forgot_password'))

        # Xóa token cũ của email này (nếu có)
        PasswordResetToken.query.filter_by(email=email).delete()

        # Tạo token mới
        reset_token = generate_reset_code()
        new_token = PasswordResetToken(email=email, token=reset_token)
        db.session.add(new_token)
        db.session.commit()

        # Gửi email chứa mã khôi phục
        send_reset_email(email, reset_token)

        return redirect(url_for('verify_code', email=email))

    return render_template('forgot_password.html')

@app.route('/verify_code', methods=['GET', 'POST'])
def verify_code():
    session.pop('_flashes', None)
    email = request.args.get('email')  # Lấy email từ query string
    if request.method == 'POST':
        reset_code = request.form['reset_code']

        # Kiểm tra token
        token = PasswordResetToken.query.filter_by(email=email, token=reset_code).first()
        if not token:
            flash('Wrong code!!!', 'warning')
            return redirect(url_for('verify_code', email=email))

        # Kiểm tra thời hạn token
        if datetime.utcnow() - token.created_at > timedelta(minutes=1):
            db.session.delete(token)
            db.session.commit()
            return redirect(url_for('forgot_password'))

        # Chuyển đến trang đặt lại mật khẩu
        return redirect(url_for('reset_password', email=email))

    return render_template('verify_code.html', email=email)

@app.route('/reset_password', methods=['GET', 'POST'])
def reset_password():
    email = request.args.get('email') or request.form.get('email')
    if request.method == 'POST':
        new_password = request.form['password']
        confirm_password = request.form['confirm_password']

        # Kiểm tra mật khẩu mới và xác nhận mật khẩu có khớp không
        if new_password != confirm_password:
            flash('Passwords do not match!', 'warning')
            return redirect(url_for('reset_password', email=email))

        # Tìm người dùng trong cơ sở dữ liệu
        user = User.query.filter_by(email=email).first()
        if not user:
            return redirect(url_for('reset_password', email=email))

        # Lưu trực tiếp mật khẩu mới vào cơ sở dữ liệu
        user.password = new_password
        db.session.commit()

        # Xóa token sau khi sử dụng
        PasswordResetToken.query.filter_by(email=email).delete()
        db.session.commit()

        flash('Password reset successfully!', 'success')
        return redirect(url_for('login'))

    return render_template('reset_password.html', email=email)

############################################################################################# Kì Thi

def delete_invalid_exams():
    """Xóa các kỳ thi đã kết thúc mà ngày kết thúc nhỏ hơn ngày hôm nay"""
    with app.app_context():  # Đảm bảo chạy trong application context
        # Lấy ngày hiện tại
        today = datetime.utcnow().date()

        # Lọc các kỳ thi đã kết thúc (ngày kết thúc < hôm nay)
        expired_exams = OrganizationExam.query.filter(db.func.date(OrganizationExam.end_time) < today).all()
        #print(f"Found {len(expired_exams)} expired exams.")  # Log số lượng kỳ thi đã hết hạn
        
        # Xóa từng kỳ thi
        for exam in expired_exams:
            print(f"Deleting expired exam with ID {exam.id} and title {exam.title}")
            db.session.delete(exam)
        
        # Xác nhận thay đổi trong cơ sở dữ liệu
        db.session.commit()

# Khởi tạo APScheduler
scheduler = BackgroundScheduler()
with app.app_context():
    scheduler.add_job(delete_invalid_exams, 'interval', seconds=10)  # Chạy mỗi ngày
    scheduler.start()

@app.route('/organization/<int:id>/create_exam', methods=['GET', 'POST'])
@login_required
def create_exam(id):
    organization = Organization.query.get_or_404(id)
    current_user_id = session.get('user_id')
    
    # Check if user is admin or manager in this organization
    user_org = UserOrganization.query.filter_by(
        user_id=current_user_id,
        organization_id=id
    ).first()
    
    if not user_org or user_org.role not in ['admin', 'manager']:
        return redirect(url_for('view_exams', id=id))
    
    if request.method == 'POST':
        title = request.form['title']
        start_time = request.form['start_time']
        end_time = request.form['end_time']
        exam_file = request.files.get('exam_file')
        answer_file = request.files.get('answer_file')

        # Validate file formats
        if not exam_file or not exam_file.filename.endswith('.pdf'):
            return redirect(request.url)
        if not answer_file or not (answer_file.filename.endswith('.doc') or answer_file.filename.endswith('.docx')):
            return redirect(request.url)

        # Save exam file
        upload_folder = os.path.join(app.config['UPLOAD_FOLDER'], 'exams')
        os.makedirs(upload_folder, exist_ok=True)
        exam_path = os.path.join(upload_folder, exam_file.filename)
        exam_file.save(exam_path)

        # Save answer file
        answer_path = os.path.join(upload_folder, answer_file.filename)
        answer_file.save(answer_path)

        # Add leading '/' to paths
        if not exam_path.startswith('/'):
            exam_path = '/' + exam_path.replace("\\", "/")
        if not answer_path.startswith('/'):
            answer_path = '/' + answer_path.replace("\\", "/")

        # Parse start and end times
        start_time = datetime.strptime(start_time, '%Y-%m-%dT%H:%M')
        end_time = datetime.strptime(end_time, '%Y-%m-%dT%H:%M')

        # Save exam to the database
        new_exam = OrganizationExam(
            title=title,
            start_time=start_time,
            end_time=end_time,
            organization_id=id,
            file_path=exam_path,  # Corrected to use exam_path
            answer_path=answer_path
        )
        db.session.add(new_exam)
        db.session.commit()

        return redirect(url_for('view_exams', id=id))

    return render_template('organization_create_exam.html', organization=organization, current_page='organization')

class UserAssignment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    exam_id = db.Column(db.Integer, db.ForeignKey('organization_exam.id'), nullable=False)
    submitted = db.Column(db.Boolean, default=False)
    submission_content = db.Column(JSON, nullable=True)  # Danh sách câu trả lời
    submitted_at = db.Column(db.DateTime, default=None)
    submission_file = db.Column(db.String(200), nullable=True)  # Đường dẫn file
    total_score = db.Column(db.Integer, nullable=True)  # Tổng điểm

@app.route('/organization/<int:id>/exams')
@login_required
def view_exams(id):
    organization = Organization.query.get_or_404(id)
    exams = OrganizationExam.query.filter_by(organization_id=id).all()
    user_id = session['user_id']

    # Lấy thông tin nộp bài cho từng kỳ thi
    submissions = {
        assignment.exam_id: assignment.submitted
        for assignment in UserAssignment.query.filter_by(user_id=user_id).all()
    }

    # Lấy thời gian hiện tại (UTC)
    now = datetime.now()

    # Lọc các kỳ thi đang diễn ra và đã qua
    ongoing_exams = [exam for exam in exams if exam.start_time <= now <= exam.end_time]
    past_exams = [exam for exam in exams if now > exam.end_time]

    return render_template(
        'organization_exam.html',
        organization=organization,
        ongoing_exams=ongoing_exams,
        past_exams=past_exams,
        submissions=submissions,
        current_page="organization"
    )

EXAM_FOLDER = os.path.join(app.config['UPLOAD_FOLDER'], 'exams')

# Thiết lập logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def compare_docx(submission_path, answer_path):
    """
    So sánh hai tài liệu docx (bài nộp và đáp án).
    Args:
        submission_path (str): Đường dẫn đến file bài nộp.
        answer_path (str): Đường dẫn đến file đáp án.
    Returns:
        dict: Kết quả so sánh với các trường correct, total, wrong, details.
    """
    result = {'correct': 0, 'total': 0, 'wrong': [], 'details': []}

    try:
        # Đọc tài liệu
        submission = Document(submission_path)
        answer = Document(answer_path)
    except Exception as e:
        logger.error(f"Lỗi khi đọc tài liệu: {str(e)}")
        return {'error': f'Không thể đọc tài liệu: {str(e)}'}

    # Hàm chuẩn hóa văn bản
    def normalize_text(text):
        # Loại bỏ khoảng trắng thừa, xuống dòng, và chuyển về lowercase
        text = re.sub(r'\s+', ' ', text.strip())
        return text.lower()

    # Lấy các đoạn văn, bỏ qua tiêu đề (heading) nếu cần
    def get_content_paragraphs(doc):
        paragraphs = []
        for para in doc.paragraphs:
            # Bỏ qua các đoạn trống hoặc tiêu đề (kiểm tra style)
            if para.text.strip() and para.style.name not in ['Heading 1', 'Heading 2', 'Heading 3']:
                paragraphs.append(normalize_text(para.text))
        return paragraphs

    # Lấy danh sách đoạn văn từ bài nộp và đáp án
    submission_text = get_content_paragraphs(submission)
    answer_text = get_content_paragraphs(answer)

    # Kiểm tra nếu tài liệu rỗng
    if not answer_text:
        logger.error("Tài liệu đáp án rỗng")
        return {'error': 'Tài liệu đáp án rỗng'}
    if not submission_text:
        logger.error("Tài liệu bài nộp rỗng")
        return {'error': 'Tài liệu bài nộp rỗng'}

    # Đếm tổng số câu hỏi
    result['total'] = len(answer_text)

    # So sánh từng đoạn văn
    for index, (ans, sub) in enumerate(zip(answer_text, submission_text), start=1):
        # Sử dụng SequenceMatcher để tính độ tương đồng
        similarity = SequenceMatcher(None, ans, sub).ratio()
        is_correct = similarity > 0.95  # Ngưỡng tương đồng (có thể điều chỉnh)

        if is_correct:
            result['correct'] += 1
            result['details'].append({
                'question': index,
                'correct': True,
                'submission': sub,
                'answer': ans,
                'similarity': round(similarity * 100, 2)
            })
        else:
            result['wrong'].append(index)
            result['details'].append({
                'question': index,
                'correct': False,
                'submission': sub,
                'answer': ans,
                'similarity': round(similarity * 100, 2)
            })

    # Xử lý trường hợp bài nộp thiếu câu trả lời
    if len(submission_text) < len(answer_text):
        for index in range(len(submission_text) + 1, len(answer_text) + 1):
            result['wrong'].append(index)
            result['details'].append({
                'question': index,
                'correct': False,
                'submission': 'No answer provided',
                'answer': normalize_text(answer_text[index - 1]),
                'similarity': 0
            })

    # Xử lý trường hợp bài nộp có thừa đoạn văn
    if len(submission_text) > len(answer_text):
        for index in range(len(answer_text) + 1, len(submission_text) + 1):
            result['details'].append({
                'question': index,
                'correct': False,
                'submission': normalize_text(submission_text[index - 1]),
                'answer': 'No corresponding answer',
                'similarity': 0
            })

    return result

@app.route('/assignment/<int:exam_id>', methods=['GET'])
@login_required
def assignment_page(exam_id):
    exam = OrganizationExam.query.get_or_404(exam_id)
    return render_template('assignment.html', exam=exam, current_page='organization')

@app.route('/submit_assignment/<int:exam_id>', methods=['POST'])
@login_required
def submit_assignment(exam_id):
    # Lấy thông tin người dùng từ session
    user_id = session['user_id']
    user = User.query.get(user_id)

    # Kiểm tra xem kỳ thi có tồn tại không
    exam = OrganizationExam.query.get(exam_id)

    # Kiểm tra xem người dùng đã nộp bài chưa
    user_assignment = UserAssignment.query.filter_by(user_id=user_id, exam_id=exam_id).first()

    # Lấy đường dẫn file đáp án từ OrganizationExam
    answer_path = os.path.join(EXAM_FOLDER, 'ANS_EXAMPLE.docx')  
    if not answer_path or not os.path.exists(answer_path):
        return jsonify({'error': 'File đáp án không tồn tại!'}), 404

    # Lấy câu trả lời từ form
    answers = request.form.getlist('answer')
    if not answers:
        return jsonify({'error': 'Không có câu trả lời nào được cung cấp!'}), 400

    # Tạo file bài nộp
    submission_path = os.path.join(EXAM_FOLDER, f'submission_{user_id}_{exam_id}.docx')
    try:
        doc = Document()
        sections = ["I", "II", "III", "IV"]
        for i, section in enumerate(sections):
            doc.add_heading(section, level=1)
            if i < len(answers) and answers[i].strip():
                doc.add_paragraph(answers[i])
            else:
                doc.add_paragraph("None")
        doc.save(submission_path)
    except Exception as e:
        return jsonify({'error': f'Lỗi khi lưu file bài nộp: {str(e)}'}), 500

    # So sánh bài nộp với đáp án
    try:
        result = compare_docx(submission_path, answer_path)
        score = ((result['correct'] * 100) // result['total']) if result['total'] > 0 else 0
    except Exception as e:
        return jsonify({'error': f'Lỗi khi so sánh bài nộp: {str(e)}'}), 500

    # Cập nhật hoặc tạo bản ghi bài nộp
    submission_date = datetime.utcnow()
    try:
        if not user_assignment:
            user_assignment = UserAssignment(
                user_id=user_id,
                exam_id=exam_id,
                submitted=True,
                submission_content=answers,  # Lưu danh sách câu trả lời dưới dạng JSON
                submission_file=submission_path,  # Lưu đường dẫn file
                total_score=int(score),  # Lưu điểm số
                submitted_at=submission_date
            )
            db.session.add(user_assignment)
        else:
            user_assignment.submitted = True
            user_assignment.submission_content = answers
            user_assignment.submission_file = submission_path
            user_assignment.total_score = int(score)
            user_assignment.submitted_at = submission_date
        db.session.commit()
    except Exception as e:
        db.session.rollback()

    # Định dạng ngày nộp bài cho giao diện
    formatted_submission_date = submission_date.strftime("%d/%m/%Y")
    length = len(answers)

    # Render template với dữ liệu
    return render_template('view_submit.html', exam=exam, 
                           date=formatted_submission_date, 
                           current_page='organization', 
                           user=user, 
                           length=length,
                           score=score,
                           answers=user_assignment.submission_content)

@app.route('/view_submit_assignment/<int:exam_id>')
@login_required
def view_submit_assignment(exam_id):
    user_id = session['user_id']
    user = User.query.get(user_id)
    exam = OrganizationExam.query.get(exam_id)

    submission_date = datetime.utcnow()
    formatted_submission_date = submission_date.strftime("%d/%m/%Y")

    assignment = UserAssignment.query.filter_by(exam_id=exam_id, user_id = user_id).first()

    answers = assignment.submission_content
    length = len(answers)
    score = assignment.total_score

    return render_template('view_submit.html', exam=exam, 
                           date=formatted_submission_date, 
                           current_page='organization', 
                           user=user, 
                           length=length,
                           score=score,
                           answers=answers)

@app.route('/helps_create_exam')
def helps_create_exam():
    return render_template('helps_create_exam.html', current_page="organization")

@app.route('/helps_do_exam')
def helps_do_exam():
    return render_template('helps_do_exam.html', current_page="organization")

#######################################################################################################################

CORS(app)

client = os.getenv("GROQ_API_KEY")

fixed_responses = {
    "python là gì": "Python là ngôn ngữ lập trình bậc cao, mạnh mẽ và dễ học.",
    "biến là gì": "Biến là vùng lưu trữ dữ liệu trong lập trình, có thể thay đổi giá trị trong quá trình chạy chương trình.",
}
@app.route("/api/chat", methods=["POST"])
def chat():
    data = request.json
    message = data.get("message", "").lower()

    if message in fixed_responses:
        return jsonify({"reply": fixed_responses[message]})

    try:
        response = client.chat.completions.create(
            model="llama3-8b-8192",
            messages=[
                {"role": "system", "content": "Bạn là trợ lý AI, luôn trả lời ngắn gọn và bằng tiếng Việt."},
                {"role": "user", "content": message}
            ]
        )
        reply = response.choices[0].message.content
    except Exception as e:
        reply = f"Lỗi khi kết nối API Groq: {str(e)}"

    return jsonify({"reply": reply})

#######################################################################################################################

class UserProgress(db.Model):
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), primary_key=True)
    overall_completion = db.Column(db.Float, default=0.0)  # Tỷ lệ hoàn thành (%)
    last_updated = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    user = db.relationship('User', backref=db.backref('progress', uselist=False))

@app.route('/thongke', methods=['GET'])
@login_required
def thongke():
    user_id = session['user_id']
    user = User.query.get(user_id)
    name = user.username
    
    # Cập nhật chuỗi tiến độ
    update_streak(user)
    
     # Tính overall_completion
    total_items = Lesson.query.count() + Challenge.query.count()
    completed_items = (
        UserLesson.query.filter_by(user_id = user.id, completed=True).count() +
        UserChallenge.query.filter_by(user_id = user.id, solved=True).count()
    )
    overall_completion = (completed_items / total_items * 0.1) if total_items > 0 else 0

    # Lưu vào UserProgress
    user_progress = UserProgress.query.filter_by(user_id = user.id).first()
    if user_progress:
        user_progress.overall_completion = overall_completion
        user_progress.last_updated = datetime.utcnow()
    else:
        user_progress = UserProgress(
            user_id = user.id,
            overall_completion = overall_completion
        )
        db.session.add(user_progress)
    db.session.commit()

    # Tính x% người dùng bị vượt qua
    total_users = User.query.count()
    users_lower = UserProgress.query.filter(
        UserProgress.overall_completion < overall_completion
    ).count()
    percentile_rank = ((users_lower) / total_users * 100) if total_users > 0 else (0)

    # Thu thập dữ liệu thống kê
    data = {
        'lessons': {
            'total': UserLesson.query.filter_by(user_id=user_id, completed_only=True).count(),
            'weekly': calculate_weekly_lessons(user_id),
            'full': Lesson.query.count()
        },
        'challenges': {
            'total': UserChallenge.query.filter_by(user_id=user_id, solved=True).count(),
            'weekly': calculate_weekly_challenges(user_id)
        },
        'daily_activity': calculate_daily_activity(user_id),
        'streak': user.streak,
        'challenge_progress': calculate_challenge_progress(user_id),
        'overall_completion': round(calculate_overall_completion(user_id), 1),
        'ai_insights': generate_ai_insights(user_id),
        'percentile_rank': percentile_rank,
    }
    
    return render_template('thongke.html', stats=data, name=name, current_page="thongke")

def calculate_weekly_lessons(user_id):
    """Tính số bài học hoàn thành trong tuần"""
    week_start = datetime.utcnow() - timedelta(days=7)
    count = UserLesson.query.filter(
        UserLesson.user_id == user_id,
        UserLesson.completed_only == True,
        UserLesson.timestamp >= week_start
    ).count()
    return count

class DailyProgress(db.Model):
    __tablename__ = 'daily_progress'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    date = db.Column(db.Date, nullable=False)
    lessons_completed = db.Column(db.Integer, default=0)
    challenges_completed = db.Column(db.Integer, default=0)

    user = db.relationship('User', backref=db.backref('daily_progress', lazy=True))

def save_daily_progress(user):
    today = datetime.utcnow().date()
    existing = DailyProgress.query.filter_by(user_id=user.id, date=today).first()
    if not existing:
        progress = DailyProgress(
            user_id=user.id,
            date=today,
            lessons_completed=user_lessons_completed_today(user.id),
            challenges_completed=user.problems_solved_today
        )
        db.session.add(progress)
    else:
        existing.lessons_completed = user_lessons_completed_today(user.id)
        existing.challenges_completed = user.problems_solved_today
    db.session.commit()

def calculate_daily_activity(user_id):
    today = datetime.utcnow().date()
    week_start = today - timedelta(days=today.weekday())  # Lấy thứ Hai

    daily_activity = []

    for i in range(7):
        day = week_start + timedelta(days=i)

        if day > today:
            # Không tính ngày tương lai
            daily_activity.append({
                'date': day.strftime('%Y-%m-%d'),
                'lessons_complete_day': 0,
                'challenges_completed': 0
            })
            continue

        if day == today:
            # Dùng bảng hiện tại cho hôm nay
            lessons_completed = UserLesson.query.filter(
                UserLesson.user_id == user_id,
                UserLesson.complete_day == True,
                func.date(UserLesson.timestamp) == day
            ).count()

            challenges_completed = UserChallenge.query.filter(
                UserChallenge.user_id == user_id,
                UserChallenge.complete_day == True,
                func.date(UserChallenge.timestamp) == day
            ).count()
        else:
            # Dùng bảng DailyProgress cho ngày trước đó
            progress = DailyProgress.query.filter_by(user_id=user_id, date=day).first()
            lessons_completed = progress.lessons_completed if progress else 0
            challenges_completed = progress.challenges_completed if progress else 0

        daily_activity.append({
            'date': day.strftime('%Y-%m-%d'),
            'lessons_complete_day': lessons_completed,
            'challenges_completed': challenges_completed
        })

    return daily_activity

def calculate_weekly_challenges(user_id):
    """Tính số thử thách hoàn thành trong tuần"""
    week_start = datetime.utcnow() - timedelta(days=7)
    count = UserChallenge.query.filter(
        UserChallenge.user_id == user_id,
        UserChallenge.solved == True,
        UserChallenge.timestamp >= week_start
    ).count()
    return count

def update_streak(user):
    """Cập nhật chuỗi tiến độ"""
    today = datetime.utcnow().date()
    if user.last_active_date:
        last_active = user.last_active_date
        if last_active == today:
            return  # Không thay đổi nếu đã hoạt động hôm nay
        elif last_active == today - timedelta(days=1):
            user.streak += 1  # Tăng chuỗi nếu hoạt động ngày hôm trước
        else:
            user.streak = 1  # Reset chuỗi nếu gián đoạn
    else:
        user.streak = 1
    user.last_active_date = today
    db.session.commit()

def calculate_challenge_progress(user_id):
    """Tính tiến độ thử thách theo cấp độ"""
    levels = {
        'beginner': {'completed': 0, 'total': Challenge.query.filter_by(difficulty='beginner').count()},
        'intermediate': {'completed': 0, 'total': Challenge.query.filter_by(difficulty='intermediate').count()},
        'advanced': {'completed': 0, 'total': Challenge.query.filter_by(difficulty='advanced').count()}
    }
    challenges = UserChallenge.query.filter_by(user_id=user_id, solved=True).join(Challenge).all()
    for uc in challenges:
        difficulty = uc.challenge.difficulty
        if difficulty in levels:
            levels[difficulty]['completed'] += 1
    return levels

def calculate_overall_completion(user_id):
    """Tính tỷ lệ hoàn thành tổng thể"""
    total_lessons = Lesson.query.count()
    total_challenges = Challenge.query.count()
    completed_lessons = UserLesson.query.filter_by(user_id=user_id, completed_only=True).count()
    completed_challenges = UserChallenge.query.filter_by(user_id=user_id, solved=True).count()
    
    total_tasks = total_lessons + total_challenges
    completed_tasks = completed_lessons + completed_challenges
    return (completed_tasks / total_tasks * 100) if total_tasks > 0 else 0

@app.route('/api/thongke', methods=['GET'])
@login_required
def api_thongke():
    user_id = session['user_id']
    user = User.query.get(user_id)
    
    # Cập nhật chuỗi tiến độ
    update_streak(user)
    
    # Thu thập dữ liệu thống kê
    data = {
        'lessons': {
            'total': UserLesson.query.filter_by(user_id=user_id, completed_only=True).count(),
            'weekly': calculate_weekly_lessons(user_id)
        },
        'challenges': {
            'total': UserChallenge.query.filter_by(user_id=user_id, solved=True).count(),
            'weekly': calculate_weekly_challenges(user_id)
        },
        'streak': user.streak,
        'challenge_progress': calculate_challenge_progress(user_id),
        'overall_completion': round(calculate_overall_completion(user_id), 1),
        'ai_insights': generate_ai_insights(user_id)
    }
    
    return jsonify(data)

def save_all_users_daily_progress():
    """Lưu tiến độ hàng ngày của tất cả người dùng vào bảng DailyProgress."""
    with app.app_context():
        yesterday = datetime.utcnow().date() - timedelta(days=1)
        users = User.query.all()

        for user in users:
            existing = DailyProgress.query.filter_by(user_id=user.id, date=yesterday).first()
            if not existing:
                lessons_completed = UserLesson.query.filter(
                    UserLesson.user_id == user.id,
                    UserLesson.complete_day == True,
                    func.date(UserLesson.timestamp) == yesterday
                ).count()

                challenges_completed = UserChallenge.query.filter(
                    UserChallenge.user_id == user.id,
                    UserChallenge.complete_day == True,
                    func.date(UserChallenge.timestamp) == yesterday
                ).count()

                progress = DailyProgress(
                    user_id=user.id,
                    date=yesterday,
                    lessons_completed=lessons_completed,
                    challenges_completed=challenges_completed
                )
                db.session.add(progress)

        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()

# Khởi tạo APScheduler
scheduler = BackgroundScheduler()
with app.app_context():
    scheduler.add_job(
        save_all_users_daily_progress,
        'cron',
        hour=0,
        minute=1,
        timezone='UTC'
    )
    scheduler.start()

def generate_ai_insights(user_id):
    today = datetime.utcnow().date()
    week_start = today - timedelta(days=6)

    # 1. Lấy dữ liệu từ bảng DailyProgress (6 ngày trước)
    daily_records = DailyProgress.query.filter(
        DailyProgress.user_id == user_id,
        DailyProgress.date >= week_start,
        DailyProgress.date < today  # không lấy hôm nay
    ).all()

    # 2. Thêm dữ liệu hôm nay từ bảng gốc
    lessons_today = UserLesson.query.filter(
        UserLesson.user_id == user_id,
        UserLesson.complete_day == True,
        func.date(UserLesson.timestamp) == today
    ).count()

    challenges_today = UserChallenge.query.filter(
        UserChallenge.user_id == user_id,
        UserChallenge.complete_day == True,
        func.date(UserChallenge.timestamp) == today
    ).count()

    # Tạo bản ghi giả cho hôm nay (không ghi vào DB)
    class Record:
        def __init__(self, date, lessons, challenges):
            self.date = date
            self.lessons_completed = lessons
            self.challenges_completed = challenges

    today_record = Record(today, lessons_today, challenges_today)
    all_records = daily_records + [today_record]

    # 3. Tổng số
    lesson_count = sum(r.lessons_completed for r in all_records)
    challenge_count = sum(r.challenges_completed for r in all_records)
    active_days = sum(1 for r in all_records if r.lessons_completed + r.challenges_completed > 0)

    # 4. Tính ngày tích cực nhất (dựa trên tổng hoạt động)
    day_names = ['Thứ Hai', 'Thứ Ba', 'Thứ Tư', 'Thứ Năm', 'Thứ Sáu', 'Thứ Bảy', 'Chủ Nhật']
    activity_by_day = Counter()
    for record in all_records:
        total = record.lessons_completed + record.challenges_completed
        if total > 0:
            day = day_names[record.date.weekday()]
            activity_by_day[day] += total

    most_active_days = [day for day, _ in activity_by_day.most_common(2)] or ['Không đủ dữ liệu']

    # 5. Chủ đề mạnh
    challenges = UserChallenge.query.filter_by(user_id=user_id, solved=True).join(Challenge).all()
    categories = [c.challenge.category for c in challenges if c.challenge.category]
    strong_categories = [cat for cat, _ in Counter(categories).most_common(1)] or ['Chưa xác định']

    # 6. Prompt AI
    prompt = f"""
    Phân tích dữ liệu học tập của người dùng:
    - Số bài học hoàn thành trong tuần: {lesson_count}
    - Số thử thách hoàn thành trong tuần: {challenge_count}
    - Số ngày hoạt động: {active_days}
    - Các ngày tích cực nhất: {', '.join(most_active_days)}
    - Chủ đề mạnh: {', '.join(strong_categories)}
    Vui lòng cung cấp phân tích bằng tiếng Việt, định dạng JSON như sau:
    ```json
    {{
        "consistency": {{
            "title": "Phân Tích Tính Liên Tục",
            "content": "Phân tích tính liên tục trong hoạt động học tập, ví dụ: ngày nào tích cực nhất và gợi ý cải thiện.",
            "tags": ["Nhận Diện Mô Hình", "Gợi Ý"]
        }},
        "strengths": {{
            "title": "Điểm Mạnh Được Xác Định",
            "content": "Điểm mạnh của người dùng, ví dụ: chủ đề nào nổi bật và cải thiện ra sao.",
            "tags": ["Hiệu Suất", "Khích Lệ"]
        }},
        "tips": {{
            "title": "Mẹo Học Tập",
            "content": "Gợi ý cá nhân hóa để cải thiện học tập, ví dụ: ôn tập chủ đề cụ thể.",
            "tags": ["Gợi Ý", "Cá Nhân Hóa"]
        }}
    }}
    ```
    Trả về chỉ JSON, không thêm text ngoài định dạng.
    """

    try:
        response = client.chat.completions.create(
            model="llama3-8b-8192",
            messages=[
                {"role": "system", "content": "Bạn là trợ lý AI, trả lời bằng JSON hợp lệ theo định dạng yêu cầu."},
                {"role": "user", "content": prompt}
            ]
        )
        reply = response.choices[0].message.content.strip()
        insights = json.loads(reply)
        return insights if all(k in insights for k in ['consistency', 'strengths', 'tips']) else {}

    except Exception as e:
        return {
            'consistency': {
                'title': 'Phân Tích Tính Liên Tục',
                'content': f'Bạn đã học {active_days} ngày trong tuần qua, tích cực nhất vào {", ".join(most_active_days)}.',
                'tags': ['Nhận Diện Mô Hình', 'Gợi Ý']
            },
            'strengths': {
                'title': 'Điểm Mạnh Được Xác Định',
                'content': f'Bạn đã hoàn thành {challenge_count} thử thách, mạnh nhất ở "{strong_categories[0]}".',
                'tags': ['Hiệu Suất', 'Khích Lệ']
            },
            'tips': {
                'title': 'Mẹo Học Tập',
                'content': f'Hãy luyện tập thêm thử thách về "{strong_categories[0]}" để nâng cao kỹ năng!',
                'tags': ['Gợi Ý', 'Cá Nhân Hóa']
            }
        }

#################################################################################################

@app.route('/contact', methods=['GET', 'POST'])
def contact():
    session.pop('_flashes', None)
    if request.method == 'POST':
        # Get form data
        name = request.form.get('name')
        email = request.form.get('email')
        subject = request.form.get('subject')
        message = request.form.get('message')

        # Validate form data
        if not all([name, email, subject, message]):
            return redirect(url_for('contact'))

        # Validate email format
        email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_regex, email):
            return redirect(url_for('contact'))

        # Prepare email
        try:
            msg = Message(
                subject=f'[CodeJourney Contact] {subject}',
                sender='kakawiner04@gmail.com',
                recipients=['kakawiner04@gmail.com']
            )
            msg.body = f"""
New Contact Message from CodeJourney Website

Name: {name}
Email: {email}
Subject: {subject}
Message: 
{message}

Organization: CodeJourney
"""
            mail.send(msg)

            flash('Your message has been sent successfully! We will get back to you soon.', 'success')
            return redirect(url_for('contact'))

        except Exception as e:
            # Log error for debugging (optional)
            flash(f'Failed to send message: {str(e)}. Please try again later.', 'warning')
            return redirect(url_for('contact'))

    return render_template('contact.html')

#################################################################################################

@app.route('/tongquan')
@login_required
def tongquan():
    user_id=session['user_id']
    user=User.query.filter_by(id=user_id).first()
    stats=UserProgress.query.filter_by(user_id=user_id).first()
    overall_completion = round(calculate_overall_completion(user_id), 1)
    return render_template('tongquan.html', current_page ='tongquan', user=user, stats=stats, overall_completion=overall_completion)

@app.route('/dgnl')
@login_required
def dgnl():
    return render_template('dgnl.html', current_page ='dgnl')

#################################################################################################

with app.app_context():
    db.create_all()

with app.app_context():
    users = User.query.all()
    o_exams = OrganizationExam.query.all()
    for o_exam in o_exams:
        if o_exam.file_path and not o_exam.file_path.startswith('/'):
            o_exam.file_path = f"/{o_exam.file_path}"  
    for user in users:
        if user.avatar and not user.avatar.startswith('/'):
            user.avatar = f"/{user.avatar}"  
    db.session.commit()

if __name__ == "__main__":
    app.run(host="0.0.0.0", port = 3000, debug=True)

#dgnl chua xong, chấm ở lesson
